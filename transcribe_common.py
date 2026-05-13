#!/usr/bin/env python3
"""
Shared helpers for the transcription CLIs.
"""
from __future__ import annotations

import argparse
import getpass
import glob
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Iterable, Optional

import requests
from docx import Document
from docx.shared import Pt
from google.auth.exceptions import RefreshError
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build

from transcript_artifacts import TranscriptArtifact
from transcript_store import ingest_artifact

DEFAULT_LANGUAGE_CODE = "en_us"
DEFAULT_OPENAI_MODEL = "gpt-4o-mini"
CALENDAR_SCOPES = ["https://www.googleapis.com/auth/calendar.readonly"]
SCRIPT_DIR = Path(__file__).resolve().parent
WILDCARD_PATTERN = re.compile(r"[*?\[\]]")
EVENT_WINDOW_BUFFER_SECONDS = 5 * 60
SENTENCE_BOUNDARY_RE = re.compile(r"(?<=[.!?])\s+")
MIN_SRT_CUE_DURATION = 0.5
LANGUAGE_CODE_ALIASES = {
    "en": "en_us",
    "en_us": "en_us",
    "en-us": "en_us",
    "en_uk": "en_uk",
    "en-uk": "en_uk",
    "en_gb": "en_uk",
    "en-gb": "en_uk",
    "en_au": "en_au",
    "en-au": "en_au",
    "english": "en_us",
    "portuguese": "pt",
    "português": "pt",
    "portugues": "pt",
    "portugese": "pt",
    "pt": "pt",
    "pt_br": "pt",
    "pt-br": "pt",
    "pt_pt": "pt",
    "pt-pt": "pt",
}


class TranscriptionError(RuntimeError):
    """Raised for CLI-level transcription and export failures."""


@dataclass
class CalendarProvider:
    name: str
    backend: Any = None
    account: Optional[str] = None
    client: Optional[str] = None
    config_dir: Optional[Path] = None
    env: Optional[dict[str, str]] = None
    credentials_path: Optional[Path] = None
    token_path: Optional[Path] = None
    fallback_client_path: Optional[Path] = None


@dataclass
class CalendarProviderConfig:
    name: str
    account: Optional[str] = None
    client: Optional[str] = None
    config_dir: Optional[Path] = None
    env: Optional[dict[str, str]] = None


DEFAULT_CALENDAR_PROVIDER_ORDER = ["gog", "gws", "google-api"]
SUPPORTED_CALENDAR_PROVIDERS = {"google-api", "google", "gog", "gws"}


@dataclass
class CalendarProviderBundle:
    providers: list[CalendarProvider]


def extract_response_detail(response: requests.Response) -> Optional[str]:
    try:
        payload = response.json()
        if isinstance(payload, dict):
            detail = payload.get("error") or payload.get("message")
            if detail:
                return str(detail)
            return json.dumps(payload, ensure_ascii=False)
        return json.dumps(payload, ensure_ascii=False)
    except ValueError:
        text = (response.text or "").strip()
        return text or None


def raise_for_status_with_details(response: requests.Response, *, context: str) -> None:
    try:
        response.raise_for_status()
    except requests.HTTPError as exc:
        detail = extract_response_detail(response)
        if detail:
            raise TranscriptionError(f"{context} failed ({response.status_code}): {detail}") from exc
        raise TranscriptionError(f"{context} failed ({response.status_code}).") from exc


def to_rfc3339(dt: datetime) -> str:
    return dt.astimezone(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def get_file_modified_time(path: Path) -> datetime:
    stats = path.stat()
    return datetime.fromtimestamp(stats.st_mtime, tz=timezone.utc)


def load_client_config(credentials_path: Path) -> dict[str, Any]:
    with credentials_path.open("r", encoding="utf-8") as handle:
        config = json.load(handle)
    if "installed" in config:
        return config
    if "web" in config:
        web_config = config["web"]
        return {
            "installed": {
                "client_id": web_config.get("client_id"),
                "project_id": web_config.get("project_id"),
                "auth_uri": web_config.get("auth_uri"),
                "token_uri": web_config.get("token_uri"),
                "auth_provider_x509_cert_url": web_config.get("auth_provider_x509_cert_url"),
                "client_secret": web_config.get("client_secret"),
                "redirect_uris": web_config.get(
                    "redirect_uris",
                    ["urn:ietf:wg:oauth:2.0:oob", "http://localhost"],
                ),
            }
        }
    raise TranscriptionError(
        f"Unsupported Google client secrets structure in {credentials_path}. Expected 'installed' or 'web'."
    )


def build_google_calendar_service(
    credentials_path: Path,
    token_path: Path,
    fallback_client_path: Optional[Path] = None,
):
    client_config: Optional[dict[str, Any]] = None
    chosen_config_path: Optional[Path] = None

    candidate_paths = [credentials_path]
    if fallback_client_path:
        candidate_paths.append(fallback_client_path)
    candidate_paths.extend(
        [
            credentials_path.with_name("client_secrets.json"),
            Path("client_secrets.json"),
            SCRIPT_DIR / "credentials.json",
            SCRIPT_DIR / "client_secrets.json",
        ]
    )

    unique_candidates: list[Path] = []
    seen_paths: set[Path] = set()
    for candidate in candidate_paths:
        if not candidate:
            continue
        resolved = candidate.expanduser().resolve()
        if resolved in seen_paths:
            continue
        seen_paths.add(resolved)
        unique_candidates.append(resolved)

    for candidate in unique_candidates:
        if not candidate.exists():
            continue
        try:
            client_config = load_client_config(candidate)
            chosen_config_path = candidate
            break
        except TranscriptionError:
            continue

    if not client_config or not chosen_config_path:
        raise TranscriptionError(
            "Google client secrets not found or invalid. Provide --calendar-credentials pointing to a "
            "client secret JSON downloaded from Google Cloud Console."
        )

    creds: Optional[Credentials] = None
    if token_path.exists():
        try:
            creds = Credentials.from_authorized_user_file(str(token_path), CALENDAR_SCOPES)
        except Exception:
            try:
                token_path.unlink(missing_ok=True)
            except OSError:
                pass
            creds = None

    credentials_updated = False
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
                credentials_updated = True
            except RefreshError:
                print("Google Calendar token expired or revoked; requesting new authorization.", file=sys.stderr)
                try:
                    token_path.unlink(missing_ok=True)
                except OSError:
                    pass
                creds = None
        if not creds or not creds.valid:
            flow = InstalledAppFlow.from_client_config(client_config, CALENDAR_SCOPES)
            creds = flow.run_local_server(port=0)
            credentials_updated = True
    if creds and credentials_updated:
        token_path.parent.mkdir(parents=True, exist_ok=True)
        token_path.write_text(creds.to_json(), encoding="utf-8")

    return build("calendar", "v3", credentials=creds, cache_discovery=False)


def build_calendar_service(
    credentials_path: Path,
    token_path: Path,
    fallback_client_path: Optional[Path] = None,
    provider_configs: Optional[list[CalendarProviderConfig]] = None,
):
    providers: list[CalendarProvider] = []
    configs = provider_configs or [CalendarProviderConfig(name=name) for name in DEFAULT_CALENDAR_PROVIDER_ORDER]

    for config in configs:
        name = canonical_calendar_provider_name(config.name)
        if name == "gog":
            if shutil.which("gog"):
                providers.append(
                    CalendarProvider(
                        name=name,
                        account=config.account,
                        client=config.client,
                        env=config.env,
                    )
                )
            continue
        if name == "gws":
            if shutil.which("gws"):
                providers.append(
                    CalendarProvider(
                        name=name,
                        config_dir=config.config_dir,
                        env=config.env,
                    )
                )
            continue
        if name == "google-api":
            providers.append(
                CalendarProvider(
                    name=name,
                    credentials_path=credentials_path,
                    token_path=token_path,
                    fallback_client_path=fallback_client_path,
                    env=config.env,
                )
            )

    if providers:
        return CalendarProviderBundle(providers=providers)

    raise TranscriptionError(
        "No calendar provider is available. Configure Google Calendar credentials or install gog/gws."
    )


def canonical_calendar_provider_name(name: str) -> str:
    normalized = str(name or "").strip().lower().replace("_", "-")
    if normalized == "google":
        return "google-api"
    if normalized not in SUPPORTED_CALENDAR_PROVIDERS:
        raise TranscriptionError(
            f"Unsupported calendar provider '{name}'. Use gog, gws, or google-api."
        )
    return normalized


def parse_calendar_provider_order(raw_value: Optional[str]) -> list[str]:
    if not raw_value:
        return list(DEFAULT_CALENDAR_PROVIDER_ORDER)
    names = [item.strip() for item in raw_value.split(",") if item.strip()]
    if not names:
        raise TranscriptionError("--calendar-providers must include at least one provider.")

    ordered: list[str] = []
    for name in names:
        canonical = canonical_calendar_provider_name(name)
        if canonical not in ordered:
            ordered.append(canonical)
    return ordered


def build_calendar_provider_configs_from_args(args: argparse.Namespace) -> list[CalendarProviderConfig]:
    configs: list[CalendarProviderConfig] = []
    for name in parse_calendar_provider_order(getattr(args, "calendar_providers", None)):
        configs.append(
            CalendarProviderConfig(
                name=name,
                account=getattr(args, "calendar_gog_account", None) if name == "gog" else None,
                client=getattr(args, "calendar_gog_client", None) if name == "gog" else None,
                config_dir=getattr(args, "calendar_gws_config_dir", None) if name == "gws" else None,
            )
        )
    return configs


def parse_event_datetime(event_time: dict[str, Any]) -> Optional[datetime]:
    if not event_time:
        return None
    raw = event_time.get("dateTime") or event_time.get("date")
    if not raw:
        return None
    if len(raw) == 10:
        raw = f"{raw}T00:00:00+00:00"
    if raw.endswith("Z"):
        raw = raw[:-1] + "+00:00"
    if "T" in raw and "+" not in raw[10:] and "-" not in raw[10:]:
        raw = f"{raw}+00:00"
    return datetime.fromisoformat(raw)


def format_event_datetime(dt_value: Optional[datetime]) -> Optional[str]:
    if not dt_value:
        return None
    return dt_value.astimezone().strftime("%Y-%m-%d %H:%M %Z")


def format_event_person(person: dict[str, Any]) -> Optional[str]:
    name = person.get("displayName")
    email = person.get("email")
    if name and email:
        return f"{name} <{email}>"
    if name:
        return name
    if email:
        return email
    return None


def extract_event_metadata(event: dict[str, Any]) -> dict[str, Any]:
    start_dt = parse_event_datetime(event.get("start", {}))
    end_dt = parse_event_datetime(event.get("end", {}))
    attendees: list[str] = []
    participants: list[str] = []
    seen_participants: set[str] = set()

    def add_participant(person: dict[str, Any], *, include_declined: bool = False) -> None:
        if not person:
            return
        if not include_declined and person.get("responseStatus") == "declined":
            return
        value = format_event_person(person)
        if not value or value in seen_participants:
            return
        participants.append(value)
        seen_participants.add(value)

    add_participant(event.get("organizer", {}), include_declined=True)
    add_participant(event.get("creator", {}), include_declined=True)

    for attendee in event.get("attendees", []):
        if attendee.get("responseStatus") == "declined":
            continue
        value = format_event_person(attendee)
        if value:
            attendees.append(value)
        add_participant(attendee)

    return {
        "summary": event.get("summary") or "Untitled Event",
        "start": start_dt,
        "end": end_dt,
        "location": event.get("location"),
        "attendees": attendees,
        "participants": participants,
        "hangoutLink": event.get("hangoutLink"),
    }


def extract_provider_error_detail(raw_output: str) -> Optional[str]:
    if not raw_output:
        return None
    try:
        payload = json.loads(raw_output)
    except json.JSONDecodeError:
        return raw_output.strip() or None

    if isinstance(payload, dict):
        error = payload.get("error")
        if isinstance(error, dict):
            message = error.get("message")
            reason = error.get("reason")
            if message and reason:
                return f"{message} ({reason})"
            if message:
                return str(message)
        message = payload.get("message")
        if message:
            return str(message)
        return json.dumps(payload, ensure_ascii=False)
    return json.dumps(payload, ensure_ascii=False)


def extract_events_from_provider_payload(payload: Any) -> list[dict[str, Any]]:
    if isinstance(payload, list):
        return [item for item in payload if isinstance(item, dict)]
    if not isinstance(payload, dict):
        return []

    for key in ("items", "events", "results", "result"):
        value = payload.get(key)
        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]

    if any(key in payload for key in ("id", "summary", "start", "end")):
        return [payload]
    return []


def extract_calendars_from_provider_payload(payload: Any) -> list[dict[str, Any]]:
    if isinstance(payload, list):
        return [item for item in payload if isinstance(item, dict)]
    if not isinstance(payload, dict):
        return []

    for key in ("items", "calendars", "results", "result"):
        value = payload.get(key)
        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]

    if any(key in payload for key in ("id", "summary", "summaryOverride")):
        return [payload]
    return []


def run_json_command(command: list[str], *, provider_name: str, env: Optional[dict[str, str]] = None) -> Any:
    command_env = os.environ.copy()
    if env:
        command_env.update(env)
    try:
        completed = subprocess.run(
            command,
            check=False,
            capture_output=True,
            text=True,
            env=command_env,
        )
    except OSError as exc:
        raise TranscriptionError(f"{provider_name} invocation failed: {exc}") from exc

    stdout = (completed.stdout or "").strip()
    stderr = (completed.stderr or "").strip()
    if completed.returncode != 0:
        detail = extract_provider_error_detail(stdout) or extract_provider_error_detail(stderr)
        if detail:
            raise TranscriptionError(f"{provider_name} failed: {detail}")
        raise TranscriptionError(f"{provider_name} failed with exit code {completed.returncode}.")

    if not stdout:
        return {}

    try:
        return json.loads(stdout)
    except json.JSONDecodeError as exc:
        raise TranscriptionError(f"{provider_name} returned invalid JSON: {exc}") from exc


def build_gog_calendar_events_command(
    calendar_id: str,
    *,
    time_min: str,
    time_max: str,
    provider: Optional[CalendarProvider] = None,
) -> list[str]:
    command = ["gog"]
    if provider and provider.account:
        command.extend(["--account", provider.account])
    if provider and provider.client:
        command.extend(["--client", provider.client])
    command.extend(
        [
            "calendar",
            "events",
            calendar_id,
            "--json",
            "--results-only",
            "--no-input",
            f"--from={time_min}",
            f"--to={time_max}",
            "--max=50",
        ]
    )
    return command


def build_gog_calendar_list_command(provider: Optional[CalendarProvider] = None) -> list[str]:
    command = ["gog"]
    if provider and provider.account:
        command.extend(["--account", provider.account])
    if provider and provider.client:
        command.extend(["--client", provider.client])
    command.extend(["calendar", "calendars", "--json", "--results-only", "--no-input"])
    return command


def build_gws_calendar_env(provider: Optional[CalendarProvider]) -> dict[str, str]:
    env: dict[str, str] = {}
    if provider and provider.env:
        env.update(provider.env)
    if provider and provider.config_dir:
        env["GOOGLE_WORKSPACE_CLI_CONFIG_DIR"] = str(provider.config_dir.expanduser())
    return env


def build_gws_calendar_events_command(calendar_id: str, *, time_min: str, time_max: str) -> list[str]:
    params = {
        "calendarId": calendar_id,
        "timeMin": time_min,
        "timeMax": time_max,
        "singleEvents": True,
        "maxResults": 50,
        "orderBy": "startTime",
    }
    return [
        "gws",
        "calendar",
        "events",
        "list",
        "--params",
        json.dumps(params, separators=(",", ":")),
        "--format",
        "json",
    ]


def build_gws_calendar_list_command() -> list[str]:
    return [
        "gws",
        "calendar",
        "calendarList",
        "list",
        "--params",
        json.dumps({"maxResults": 250}, separators=(",", ":")),
        "--format",
        "json",
    ]


def list_events_via_google_api(
    service: Any,
    calendar_id: str,
    *,
    time_min: str,
    time_max: str,
) -> list[dict[str, Any]]:
    result = (
        service.events()
        .list(
            calendarId=calendar_id,
            timeMin=time_min,
            timeMax=time_max,
            singleEvents=True,
            maxResults=50,
            orderBy="startTime",
        )
        .execute()
    )
    return extract_events_from_provider_payload(result)


def list_calendars_via_google_api(service: Any) -> list[dict[str, Any]]:
    calendars: list[dict[str, Any]] = []
    page_token: Optional[str] = None
    while True:
        request = service.calendarList().list(maxResults=250)
        if page_token:
            request = service.calendarList().list(maxResults=250, pageToken=page_token)
        result = request.execute()
        calendars.extend(extract_calendars_from_provider_payload(result))
        page_token = result.get("nextPageToken") if isinstance(result, dict) else None
        if not page_token:
            break
    return calendars


def list_events_via_gog(
    calendar_id: str,
    *,
    time_min: str,
    time_max: str,
    provider: Optional[CalendarProvider] = None,
) -> list[dict[str, Any]]:
    payload = run_json_command(
        build_gog_calendar_events_command(
            calendar_id,
            time_min=time_min,
            time_max=time_max,
            provider=provider,
        ),
        provider_name="gog calendar events",
        env=provider.env if provider else None,
    )
    return extract_events_from_provider_payload(payload)


def list_calendars_via_gog(provider: Optional[CalendarProvider] = None) -> list[dict[str, Any]]:
    payload = run_json_command(
        build_gog_calendar_list_command(provider),
        provider_name="gog calendar calendars",
        env=provider.env if provider else None,
    )
    return extract_calendars_from_provider_payload(payload)


def list_events_via_gws(
    calendar_id: str,
    *,
    time_min: str,
    time_max: str,
    provider: Optional[CalendarProvider] = None,
) -> list[dict[str, Any]]:
    payload = run_json_command(
        build_gws_calendar_events_command(calendar_id, time_min=time_min, time_max=time_max),
        provider_name="gws calendar events list",
        env=build_gws_calendar_env(provider),
    )
    return extract_events_from_provider_payload(payload)


def list_calendars_via_gws(provider: Optional[CalendarProvider] = None) -> list[dict[str, Any]]:
    payload = run_json_command(
        build_gws_calendar_list_command(),
        provider_name="gws calendar calendarList list",
        env=build_gws_calendar_env(provider),
    )
    return extract_calendars_from_provider_payload(payload)


def normalize_calendar_providers(service: Any) -> list[CalendarProvider]:
    if isinstance(service, CalendarProviderBundle):
        return service.providers
    return [CalendarProvider(name="google-api", backend=service)]


def ensure_google_calendar_provider_backend(provider: CalendarProvider) -> None:
    if provider.backend is not None:
        return
    if not provider.credentials_path or not provider.token_path:
        raise TranscriptionError("Google Calendar credentials/token paths are not configured.")
    provider.backend = build_google_calendar_service(
        provider.credentials_path,
        provider.token_path,
        fallback_client_path=provider.fallback_client_path,
    )


def calendar_display_name(calendar: dict[str, Any], calendar_id: str) -> str:
    return str(
        calendar.get("summaryOverride")
        or calendar.get("summary")
        or calendar.get("name")
        or calendar_id
    )


def normalize_calendar_entry(calendar: dict[str, Any]) -> Optional[dict[str, Any]]:
    calendar_id = calendar.get("id") or calendar.get("calendarId") or calendar.get("CalendarID")
    if not calendar_id:
        return None
    calendar_id = str(calendar_id)
    return {
        "id": calendar_id,
        "summary": calendar_display_name(calendar, calendar_id),
        "accessRole": calendar.get("accessRole"),
        "primary": calendar.get("primary"),
        "timeZone": calendar.get("timeZone"),
    }


def list_calendars_for_provider(provider: CalendarProvider) -> list[dict[str, Any]]:
    if provider.name == "google-api":
        ensure_google_calendar_provider_backend(provider)
        raw_calendars = list_calendars_via_google_api(provider.backend)
    elif provider.name == "gog":
        raw_calendars = list_calendars_via_gog(provider)
    elif provider.name == "gws":
        raw_calendars = list_calendars_via_gws(provider)
    else:
        return []

    calendars: list[dict[str, Any]] = []
    seen: set[str] = set()
    for raw_calendar in raw_calendars:
        calendar = normalize_calendar_entry(raw_calendar)
        if not calendar or calendar["id"] in seen:
            continue
        calendars.append(calendar)
        seen.add(calendar["id"])
    return calendars


def list_events_for_provider(
    provider: CalendarProvider,
    calendar_id: str,
    *,
    time_min: str,
    time_max: str,
) -> list[dict[str, Any]]:
    if provider.name == "google-api":
        ensure_google_calendar_provider_backend(provider)
        return list_events_via_google_api(
            provider.backend,
            calendar_id,
            time_min=time_min,
            time_max=time_max,
        )
    if provider.name == "gog":
        return list_events_via_gog(
            calendar_id,
            time_min=time_min,
            time_max=time_max,
            provider=provider,
        )
    if provider.name == "gws":
        return list_events_via_gws(
            calendar_id,
            time_min=time_min,
            time_max=time_max,
            provider=provider,
        )
    raise TranscriptionError(f"unsupported provider {provider.name}")


def describe_matching_calendars(
    matching_events: list[dict[str, Any]],
    calendar: dict[str, Any],
) -> list[dict[str, Any]]:
    descriptions: list[dict[str, Any]] = []
    for match in matching_events:
        event = match.get("event") or {}
        descriptions.append(
            {
                "calendar_id": calendar["id"],
                "calendar_summary": calendar["summary"],
                "accessRole": calendar.get("accessRole"),
                "event_id": event.get("id"),
                "event_summary": event.get("summary") or "Untitled Event",
                "event_start": match.get("start"),
                "event_end": match.get("end"),
                "overlap_seconds": match.get("overlap_seconds"),
                "coverage": match.get("coverage"),
            }
        )
    return descriptions


def find_matching_calendars_for_provider(
    provider: CalendarProvider,
    *,
    requested_calendar_id: str,
    recording_start: datetime,
    recording_end: datetime,
    time_min: str,
    time_max: str,
) -> list[dict[str, Any]]:
    matching_calendars: list[dict[str, Any]] = []
    seen_pairs: set[tuple[str, str]] = set()

    try:
        calendars = list_calendars_for_provider(provider)
    except Exception as exc:
        print(f"Calendar lookup: failed to list calendars for provider {provider.name} ({exc}).", file=sys.stderr)
        return []

    if not any(calendar["id"] == requested_calendar_id for calendar in calendars):
        calendars.insert(
            0,
            {
                "id": requested_calendar_id,
                "summary": requested_calendar_id,
                "accessRole": None,
                "primary": requested_calendar_id == "primary",
                "timeZone": None,
            },
        )

    for calendar in calendars:
        try:
            events = list_events_for_provider(
                provider,
                calendar["id"],
                time_min=time_min,
                time_max=time_max,
            )
        except Exception as exc:
            print(
                f"Calendar lookup: provider {provider.name} failed calendar {calendar['id']} ({exc}).",
                file=sys.stderr,
            )
            continue

        calendar_matches, _ = score_calendar_events(
            events,
            recording_start=recording_start,
            recording_end=recording_end,
        )
        for description in describe_matching_calendars(calendar_matches, calendar):
            pair = (description["calendar_id"], str(description.get("event_id") or description.get("event_summary")))
            if pair in seen_pairs:
                continue
            matching_calendars.append(description)
            seen_pairs.add(pair)

    matching_calendars.sort(
        key=lambda item: (
            str(item.get("calendar_summary") or "").lower(),
            str(item.get("event_start") or ""),
            str(item.get("event_summary") or "").lower(),
        )
    )
    return matching_calendars


def attach_matching_calendars(event_info: dict[str, Any], matching_calendars: list[dict[str, Any]]) -> dict[str, Any]:
    event_info = dict(event_info)
    event_info["matching_calendars"] = matching_calendars
    return event_info


def ensure_selected_calendar_context(
    *,
    calendar_id: str,
    matching_events: list[dict[str, Any]],
    best_event: Optional[dict[str, Any]],
    matching_calendars: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    if matching_calendars:
        return matching_calendars

    selected_event: Optional[dict[str, Any]] = None
    selected_match: Optional[dict[str, Any]] = None
    if matching_events:
        selected_match = matching_events[0]
        selected_event = selected_match.get("event")
    elif best_event:
        selected_event = best_event

    if not selected_event:
        return []

    return [
        {
            "calendar_id": calendar_id,
            "calendar_summary": calendar_id,
            "accessRole": None,
            "event_id": selected_event.get("id"),
            "event_summary": selected_event.get("summary") or "Untitled Event",
            "event_start": selected_match.get("start") if selected_match else parse_event_datetime(selected_event.get("start", {})),
            "event_end": selected_match.get("end") if selected_match else parse_event_datetime(selected_event.get("end", {})),
            "overlap_seconds": selected_match.get("overlap_seconds") if selected_match else None,
            "coverage": selected_match.get("coverage") if selected_match else None,
        }
    ]


def score_calendar_events(
    events: list[dict[str, Any]],
    *,
    recording_start: datetime,
    recording_end: datetime,
) -> tuple[list[dict[str, Any]], Optional[dict[str, Any]]]:
    midpoint = recording_start + (recording_end - recording_start) / 2
    best_event: Optional[dict[str, Any]] = None
    best_score: Optional[tuple[Any, ...]] = None
    matching_events: list[dict[str, Any]] = []

    for event in events:
        event_start = parse_event_datetime(event.get("start", {}))
        event_end = parse_event_datetime(event.get("end", {}))
        if not event_start:
            continue

        event_range_end = event_end or event_start
        if event_range_end < event_start:
            event_range_end = event_start

        overlap_start = max(recording_start, event_start)
        overlap_end = min(recording_end, event_range_end)
        overlap_seconds = max((overlap_end - overlap_start).total_seconds(), 0.0)

        event_duration_seconds = max((event_range_end - event_start).total_seconds(), 0.0)
        coverage = overlap_seconds / event_duration_seconds if event_duration_seconds > 0 else 0.0

        if coverage >= 0.5:
            matching_events.append(
                {
                    "event": event,
                    "start": event_start,
                    "end": event_range_end,
                    "overlap_seconds": overlap_seconds,
                    "coverage": coverage,
                }
            )

        distances = [abs((event_start - midpoint).total_seconds())]
        if event_end:
            distances.append(abs((event_end - midpoint).total_seconds()))
        distance = min(distances)
        if overlap_seconds > 0:
            score: tuple[Any, ...] = (0, -overlap_seconds, distance)
        else:
            score = (1, distance)

        if best_score is None or score < best_score:
            best_score = score
            best_event = event

    matching_events.sort(key=lambda item: (item["start"], -item["overlap_seconds"], -item["coverage"]))
    return matching_events, best_event


def find_matching_events(
    service,
    calendar_id: str,
    recording_start: datetime,
    recording_end: datetime,
    window_hours: float,
) -> tuple[list[dict[str, Any]], Optional[dict[str, Any]], list[dict[str, Any]]]:
    if recording_end < recording_start:
        recording_end = recording_start

    time_min = to_rfc3339(recording_start - timedelta(hours=window_hours))
    time_max = to_rfc3339(recording_end + timedelta(hours=window_hours))
    provider_errors: list[str] = []
    successful_lookup = False

    for provider in normalize_calendar_providers(service):
        try:
            print(f"Calendar lookup: trying provider {provider.name}...", file=sys.stderr)
            events = list_events_for_provider(provider, calendar_id, time_min=time_min, time_max=time_max)
        except Exception as exc:
            provider_errors.append(f"{provider.name}: {exc}")
            print(f"Calendar lookup: provider {provider.name} failed ({exc}).", file=sys.stderr)
            continue

        successful_lookup = True
        print(f"Calendar lookup: provider {provider.name} returned {len(events)} event(s).", file=sys.stderr)
        matching_calendars = find_matching_calendars_for_provider(
            provider,
            requested_calendar_id=calendar_id,
            recording_start=recording_start,
            recording_end=recording_end,
            time_min=time_min,
            time_max=time_max,
        )
        if matching_calendars:
            print(
                f"Calendar lookup: provider {provider.name} found "
                f"{len(matching_calendars)} matching calendar event(s).",
                file=sys.stderr,
            )
        if not events:
            continue
        matching_events, best_event = score_calendar_events(
            events,
            recording_start=recording_start,
            recording_end=recording_end,
        )
        matching_calendars = ensure_selected_calendar_context(
            calendar_id=calendar_id,
            matching_events=matching_events,
            best_event=best_event,
            matching_calendars=matching_calendars,
        )
        return matching_events, best_event, matching_calendars

    if successful_lookup:
        return [], None, []
    if provider_errors:
        raise TranscriptionError("all calendar providers failed: " + "; ".join(provider_errors))
    return [], None, []


def sanitize_filename_part(value: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]", " ", value)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def unique_path(base_path: Path) -> Path:
    counter = 1
    candidate = base_path
    while candidate.exists():
        candidate = base_path.with_name(f"{base_path.stem} ({counter}){base_path.suffix}")
        counter += 1
    return candidate


def build_event_base_name(recording_time: datetime, event_summary: str, source_stem: str) -> str:
    timestamp_str = recording_time.astimezone().strftime("%Y-%m-%d %H-%M")
    cleaned_summary = sanitize_filename_part(event_summary or "Untitled Event")
    cleaned_source = sanitize_filename_part(source_stem)
    parts = [timestamp_str, cleaned_summary, cleaned_source]
    return " ".join(part for part in parts if part)


def rename_audio_with_event(
    audio_path: Path,
    event_info: dict[str, Any],
    recording_time: datetime,
    *,
    summary_suffix: str = "",
) -> tuple[Path, str]:
    event_summary = event_info.get("summary", "Untitled Event")
    if summary_suffix:
        event_summary = f"{event_summary} {summary_suffix}"
    base_name = build_event_base_name(recording_time, event_summary, audio_path.stem)
    target_path = unique_path(audio_path.with_name(f"{base_name}{audio_path.suffix}"))
    if target_path != audio_path:
        audio_path = audio_path.rename(target_path)
    return audio_path, base_name


def compute_event_window(
    recording_start: datetime,
    recording_end: datetime,
    event_start: Optional[datetime],
    event_end: Optional[datetime],
    duration_seconds: float,
    buffer_seconds: float,
) -> tuple[float, float]:
    event_start = event_start or event_end or recording_start
    event_end = event_end or event_start

    window_start_dt = max(recording_start, event_start - timedelta(seconds=buffer_seconds))
    window_end_dt = min(recording_end, event_end + timedelta(seconds=buffer_seconds))

    if window_end_dt < window_start_dt:
        window_end_dt = window_start_dt

    start_offset = (window_start_dt - recording_start).total_seconds()
    end_offset = (window_end_dt - recording_start).total_seconds()
    total_duration = max(duration_seconds, 0.0)

    start_offset = min(max(0.0, start_offset), total_duration)
    end_offset = min(max(0.0, end_offset), total_duration)
    if end_offset < start_offset:
        end_offset = start_offset

    return start_offset, end_offset


def select_utterances_for_window(
    utterances: list[dict[str, Any]],
    start_seconds: float,
    end_seconds: float,
) -> list[dict[str, Any]]:
    if end_seconds <= start_seconds:
        return [dict(utterance) for utterance in utterances]

    selected: list[dict[str, Any]] = []
    for utterance in utterances:
        utter_start = (utterance.get("start") or 0) / 1000
        utter_end = (utterance.get("end") or utter_start) / 1000
        if utter_end < start_seconds or utter_start > end_seconds:
            continue
        selected.append(dict(utterance))

    if not selected:
        return [dict(utterance) for utterance in utterances]
    return selected


def expand_audio_inputs(inputs: Iterable[str]) -> list[Path]:
    resolved: list[Path] = []
    seen: set[Path] = set()
    for raw in inputs:
        expanded = Path(raw).expanduser()
        if WILDCARD_PATTERN.search(str(raw)):
            matches = glob.glob(str(expanded), recursive=True)
        else:
            if expanded.exists():
                matches = [str(expanded)]
            else:
                matches = glob.glob(str(expanded), recursive=True)

        if not matches:
            print(f"Warning: no files matched '{raw}'", file=sys.stderr)
            continue

        for match in matches:
            match_path = Path(match).expanduser().resolve()
            if not match_path.is_file():
                print(f"Skipping non-file match: {match_path}", file=sys.stderr)
                continue
            if match_path in seen:
                continue
            seen.add(match_path)
            resolved.append(match_path)
    return resolved


def resolve_config_candidates(configured_path: Path) -> list[Path]:
    if configured_path.is_absolute():
        return [configured_path]
    candidates = [
        (SCRIPT_DIR / configured_path).resolve(),
        (SCRIPT_DIR / configured_path.name).resolve(),
        (Path.cwd() / configured_path).resolve(),
    ]
    unique: list[Path] = []
    seen: set[Path] = set()
    for candidate in candidates:
        if candidate in seen:
            continue
        seen.add(candidate)
        unique.append(candidate)
    return unique


def load_json_config(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        payload = json.load(handle)
    if not isinstance(payload, dict):
        raise TranscriptionError(f"Invalid JSON structure in {path}; expected an object.")
    return payload


def normalize_language_code(raw_value: Optional[str]) -> tuple[Optional[str], bool]:
    if raw_value is None:
        return DEFAULT_LANGUAGE_CODE, False
    normalized = raw_value.strip()
    if not normalized:
        return DEFAULT_LANGUAGE_CODE, False
    lowered = normalized.lower()
    if lowered in {"auto", "detect", "language_detection", "language-detection"}:
        return None, True
    return LANGUAGE_CODE_ALIASES.get(lowered, normalized), False


def resolve_language_settings(args: argparse.Namespace) -> tuple[Optional[str], bool]:
    if getattr(args, "language", None):
        return normalize_language_code(args.language)

    for candidate in resolve_config_candidates(Path(args.api_key_file).expanduser()):
        if not candidate.exists():
            continue
        try:
            payload = load_json_config(candidate)
        except json.JSONDecodeError as exc:
            print(f"Warning: skipping invalid JSON in {candidate} ({exc}); using English defaults.", file=sys.stderr)
            return DEFAULT_LANGUAGE_CODE, False
        except TranscriptionError as exc:
            print(f"Warning: skipping {candidate} ({exc}); using English defaults.", file=sys.stderr)
            return DEFAULT_LANGUAGE_CODE, False

        for field in ("assemblyai_language_code", "assembly_ai_language_code"):
            if field in payload:
                return normalize_language_code(payload.get(field))
        break

    return DEFAULT_LANGUAGE_CODE, False


def resolve_openai_key(args: argparse.Namespace) -> tuple[Optional[str], Optional[str]]:
    direct_key = getattr(args, "openai_api_key", None)
    if direct_key:
        return direct_key, "--openai-api-key"

    if getattr(args, "openai_api_key_prompt", False):
        if not sys.stdin.isatty():
            raise TranscriptionError("--openai-api-key-prompt requires an interactive terminal (stdin is not a TTY).")
        entered = getpass.getpass("OpenAI API key: ").strip()
        if entered:
            return entered, "--openai-api-key-prompt"

    if getattr(args, "openai_api_key_stdin", False):
        entered = sys.stdin.readline().strip() if sys.stdin.isatty() else sys.stdin.read().strip()
        if entered:
            return entered, "--openai-api-key-stdin"

    if not getattr(args, "api_key_file", None):
        env_key = os.getenv("OPENAI_API_KEY")
        if env_key:
            return env_key, "OPENAI_API_KEY"
        return None, None

    for candidate in resolve_config_candidates(Path(args.api_key_file).expanduser()):
        if not candidate.exists():
            continue
        try:
            payload = load_json_config(candidate)
        except json.JSONDecodeError as exc:
            print(f"Warning: skipping invalid JSON in {candidate} ({exc}).", file=sys.stderr)
            continue
        except TranscriptionError as exc:
            print(f"Warning: skipping {candidate} ({exc}).", file=sys.stderr)
            continue

        for key_field in ("openai_api_key", "open_ai_api_key"):
            openai_key = payload.get(key_field)
            if openai_key:
                return openai_key, str(candidate)
        break

    env_key = os.getenv("OPENAI_API_KEY")
    if env_key:
        return env_key, "OPENAI_API_KEY"
    return None, None


def split_into_sentences(text: str) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []
    sentences = SENTENCE_BOUNDARY_RE.split(text)
    return [sentence.strip() for sentence in sentences if sentence.strip()]


def openai_translate_lines(
    *,
    api_key: str,
    model: str,
    target_language: str,
    lines: list[str],
    timeout: float = 90.0,
) -> list[str]:
    endpoint = "https://api.openai.com/v1/chat/completions"
    session = requests.Session()
    session.headers.update({"authorization": f"Bearer {api_key}"})

    system_prompt = (
        "You are a translation engine. Translate the provided JSON array of strings into "
        f"{target_language}. Output ONLY a JSON array of strings of the same length, in the same order. "
        "Do not add commentary, numbering, or extra keys."
    )
    payload = {
        "model": model,
        "temperature": 0,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps(lines, ensure_ascii=False)},
        ],
    }
    response = session.post(endpoint, json=payload, timeout=timeout)
    raise_for_status_with_details(response, context="OpenAI translation")
    data = response.json()
    content = (((data.get("choices") or [{}])[0]).get("message") or {}).get("content") or ""
    content = content.strip()
    if not content:
        raise TranscriptionError("OpenAI translation returned an empty response.")

    try:
        translated = json.loads(content)
    except json.JSONDecodeError as exc:
        raise TranscriptionError("OpenAI translation did not return valid JSON.") from exc

    if not isinstance(translated, list) or not all(isinstance(item, str) for item in translated):
        raise TranscriptionError("OpenAI translation returned unexpected JSON (expected a string array).")
    if len(translated) != len(lines):
        raise TranscriptionError(f"OpenAI translation returned {len(translated)} lines but expected {len(lines)}.")
    return [item.strip() for item in translated]


def translate_utterances_openai(
    utterances: list[dict[str, Any]],
    *,
    api_key: str,
    model: str,
    target_language: str,
) -> list[dict[str, Any]]:
    translated_utterances = [dict(utterance) for utterance in utterances]
    sentence_map: list[list[str]] = []
    all_sentences: list[str] = []

    for utterance in translated_utterances:
        sentences = split_into_sentences(utterance.get("text") or "")
        if not sentences:
            sentence_map.append([])
            continue
        sentence_map.append(sentences)
        all_sentences.extend(sentences)

    if not all_sentences:
        return translated_utterances

    translated_sentences: list[str] = []
    batch_size = 25
    for start in range(0, len(all_sentences), batch_size):
        batch = all_sentences[start : start + batch_size]
        translated_sentences.extend(
            openai_translate_lines(
                api_key=api_key,
                model=model,
                target_language=target_language,
                lines=batch,
            )
        )

    idx = 0
    for utterance, original_sentences in zip(translated_utterances, sentence_map):
        if not original_sentences:
            utterance.pop("sentences", None)
            continue
        translated_for_utterance = translated_sentences[idx : idx + len(original_sentences)]
        idx += len(original_sentences)
        utterance["sentences"] = translated_for_utterance
        utterance["text"] = " ".join(translated_for_utterance).strip()

    return translated_utterances


def format_srt_timestamp(seconds: float) -> str:
    seconds = max(0.0, seconds)
    total_ms = int(round(seconds * 1000))
    hours, remainder = divmod(total_ms, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, millis = divmod(remainder, 1_000)
    return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"


def allocate_sentence_timings(
    start: float,
    end: float,
    sentences: list[str],
) -> list[tuple[float, float]]:
    if not sentences:
        return []
    start = max(0.0, start)
    duration = max(0.0, end - start)
    if duration <= 0.0:
        duration = MIN_SRT_CUE_DURATION * len(sentences)
        end = start + duration
    else:
        end = start + duration

    weights = [max(len(sentence.split()), 1) for sentence in sentences]
    total_weight = sum(weights) or len(sentences)

    timings: list[tuple[float, float]] = []
    current_start = start
    for idx, weight in enumerate(weights):
        if idx == len(weights) - 1:
            cue_end = max(current_start + MIN_SRT_CUE_DURATION, end)
        else:
            share = weight / total_weight
            cue_duration = duration * share
            cue_end = current_start + cue_duration
            if cue_end - current_start < MIN_SRT_CUE_DURATION:
                cue_end = current_start + MIN_SRT_CUE_DURATION
        timings.append((current_start, cue_end))
        current_start = cue_end
    return timings


def write_srt(
    utterances: list[dict[str, Any]],
    output_path: Path,
    *,
    suppress_speaker: bool,
) -> None:
    cues: list[tuple[int, float, float, str]] = []
    cue_index = 1
    previous_end = 0.0
    for utterance in utterances:
        start = (utterance.get("start") or 0) / 1000.0
        end = (utterance.get("end") or 0) / 1000.0
        if start < previous_end:
            start = previous_end
        if end < start:
            end = start
        sentences = utterance.get("sentences")
        if not isinstance(sentences, list) or not all(isinstance(item, str) for item in sentences):
            sentences = split_into_sentences(utterance.get("text") or "")
        if not sentences:
            continue
        timings = allocate_sentence_timings(start, end, sentences)
        speaker = (utterance.get("speaker") or "").strip() or "Speaker"
        for sentence, (sentence_start, sentence_end) in zip(sentences, timings):
            if sentence_end <= sentence_start:
                sentence_end = sentence_start + MIN_SRT_CUE_DURATION
            text = sentence if suppress_speaker else f"{speaker}: {sentence}"
            cues.append((cue_index, sentence_start, sentence_end, text))
            cue_index += 1
            previous_end = max(previous_end, sentence_end)

    with output_path.open("w", encoding="utf-8") as handle:
        for idx, start_ts, end_ts, text in cues:
            handle.write(f"{idx}\n")
            handle.write(f"{format_srt_timestamp(start_ts)} --> {format_srt_timestamp(end_ts)}\n")
            handle.write(f"{text}\n\n")


def determine_subtitle_codec(media_suffix: str) -> str:
    suffix = media_suffix.lower()
    if suffix in {".mp4", ".m4v", ".mov"}:
        return "mov_text"
    if suffix == ".mkv":
        return "srt"
    raise TranscriptionError(
        f"Embedding subtitles is only supported for MP4/M4V/MOV/MKV files; unsupported suffix '{media_suffix}'."
    )


def embed_subtitles_with_ffmpeg(source_media: Path, subtitle_file: Path, output_media: Path) -> None:
    if not source_media.exists():
        raise TranscriptionError(f"Source media file not found: {source_media}")
    if not subtitle_file.exists():
        raise TranscriptionError(f"Subtitle file not found: {subtitle_file}")

    subtitle_codec = determine_subtitle_codec(source_media.suffix)
    command = [
        "ffmpeg",
        "-y",
        "-i",
        str(source_media),
        "-i",
        str(subtitle_file),
        "-c",
        "copy",
        "-c:s",
        subtitle_codec,
        "-map",
        "0",
        "-map",
        "1",
        str(output_media),
    ]
    try:
        completed = subprocess.run(command, check=False, capture_output=True, text=True)
    except FileNotFoundError as exc:
        raise TranscriptionError("ffmpeg executable not found on PATH; install ffmpeg to embed subtitles.") from exc

    if completed.returncode != 0:
        stderr = (completed.stderr or "").strip()
        raise TranscriptionError(f"ffmpeg failed to embed subtitles ({stderr or 'no error output provided'})")


def format_utterance(utterance: dict[str, Any]) -> str:
    speaker = utterance.get("speaker") or "Speaker"
    start = (utterance.get("start") or 0) / 1000
    end = (utterance.get("end") or 0) / 1000
    text = (utterance.get("text") or "").strip()
    return f"{speaker} [{start:0.2f}s - {end:0.2f}s]: {text}"


def add_event_metadata_to_doc(document: Document, event_info: dict[str, Any]) -> None:
    document.add_heading("Event Details", level=2)
    document.add_paragraph(f"Event: {event_info['summary']}")

    start_str = format_event_datetime(event_info.get("start"))
    end_str = format_event_datetime(event_info.get("end"))
    if start_str:
        document.add_paragraph(f"Start: {start_str}")
    if end_str:
        document.add_paragraph(f"End: {end_str}")
    if event_info.get("location"):
        document.add_paragraph(f"Location: {event_info['location']}")
    if event_info.get("hangoutLink"):
        document.add_paragraph(f"Meeting Link: {event_info['hangoutLink']}")

    participants = event_info.get("participants") or event_info.get("attendees") or []
    if participants:
        document.add_paragraph("Participants:")
        for participant in participants:
            document.add_paragraph(participant, style="List Bullet")
    document.add_paragraph("")


def write_docx(
    utterances: list[dict[str, Any]],
    output_path: Path,
    *,
    event_info: Optional[dict[str, Any]] = None,
    title: str = "AssemblyAI Transcript",
) -> None:
    document = Document()
    document.add_heading(title, level=1)
    if event_info:
        add_event_metadata_to_doc(document, event_info)
    for utterance in utterances:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(format_utterance(utterance))
        run.font.size = Pt(11)
    document.save(output_path)


def write_text(
    utterances: list[dict[str, Any]],
    output_path: Path,
    *,
    event_info: Optional[dict[str, Any]] = None,
) -> None:
    with output_path.open("w", encoding="utf-8") as handle:
        if event_info:
            handle.write(f"Event: {event_info['summary']}\n")
            start_str = format_event_datetime(event_info.get("start"))
            end_str = format_event_datetime(event_info.get("end"))
            if start_str:
                handle.write(f"Start: {start_str}\n")
            if end_str:
                handle.write(f"End: {end_str}\n")
            if event_info.get("location"):
                handle.write(f"Location: {event_info['location']}\n")
            if event_info.get("hangoutLink"):
                handle.write(f"Meeting Link: {event_info['hangoutLink']}\n")
            participants = event_info.get("participants") or event_info.get("attendees") or []
            if participants:
                handle.write("Participants:\n")
                for participant in participants:
                    handle.write(f" - {participant}\n")
            handle.write("\n")
        for utterance in utterances:
            handle.write(format_utterance(utterance))
            handle.write("\n")


def ensure_output_dir(path: Optional[Path], audio_path: Path) -> Path:
    if path is None:
        return audio_path.parent
    path.mkdir(parents=True, exist_ok=True)
    return path


def derive_duration_seconds(
    utterances: list[dict[str, Any]],
    reported_duration_seconds: Optional[float],
) -> float:
    if reported_duration_seconds is not None:
        try:
            duration_seconds = float(reported_duration_seconds)
        except (TypeError, ValueError):
            duration_seconds = 0.0
    else:
        duration_seconds = 0.0

    if duration_seconds <= 0.0:
        max_end_ms = max((utterance.get("end") or 0) for utterance in utterances) if utterances else 0
        duration_seconds = max_end_ms / 1000 if max_end_ms else 0.0

    if duration_seconds < 0.0:
        duration_seconds = 0.0
    return duration_seconds


def process_transcription_outputs(
    audio_path: Path,
    utterances: list[dict[str, Any]],
    reported_duration_seconds: Optional[float],
    args: argparse.Namespace,
    calendar_service,
    *,
    docx_title: str,
    backend_name: str = "unknown",
) -> bool:
    if args.translate_to:
        try:
            openai_key, openai_source = resolve_openai_key(args)
        except TranscriptionError as exc:
            print(f"Translation setup failed: {exc}", file=sys.stderr)
            return False
        if not openai_key:
            print(
                "Error: --translate-to requires an OpenAI key. Provide --openai-api-key, set OPENAI_API_KEY, "
                "or add openai_api_key to api_keys.json.",
                file=sys.stderr,
            )
            return False
        if getattr(args, "print_key_sources", False) and openai_source:
            print(f"OpenAI API key source: {openai_source}", file=sys.stderr, flush=True)
        print(f"Translating transcript to {args.translate_to}...", flush=True)
        try:
            utterances = translate_utterances_openai(
                utterances,
                api_key=openai_key,
                model=args.openai_model,
                target_language=args.translate_to,
            )
        except (requests.RequestException, TranscriptionError) as exc:
            print(f"Translation failed: {exc}", file=sys.stderr)
            return False

    duration_seconds = derive_duration_seconds(utterances, reported_duration_seconds)
    working_path = audio_path

    recording_end = get_file_modified_time(working_path)
    recording_start = recording_end - timedelta(seconds=duration_seconds)
    if recording_start > recording_end:
        recording_start = recording_end

    source_stem = audio_path.stem
    primary_event_info: Optional[dict[str, Any]] = None
    transcript_jobs: list[dict[str, Any]] = []

    if args.use_calendar:
        if calendar_service is None:
            print("Warning: calendar service unavailable; skipping event lookup.", file=sys.stderr)
        else:
            try:
                matching_events, fallback_event, matching_calendars = find_matching_events(
                    calendar_service,
                    args.calendar_id,
                    recording_start,
                    recording_end,
                    args.calendar_window,
                )

                if matching_events:
                    for idx, match in enumerate(matching_events):
                        info = attach_matching_calendars(
                            extract_event_metadata(match["event"]),
                            matching_calendars,
                        )
                        event_start = match.get("start") or recording_start
                        event_end = match.get("end") or event_start
                        window_start, window_end = compute_event_window(
                            recording_start,
                            recording_end,
                            event_start,
                            event_end,
                            duration_seconds,
                            EVENT_WINDOW_BUFFER_SECONDS,
                        )
                        base_name = build_event_base_name(
                            event_start,
                            info.get("summary", "Untitled Event"),
                            source_stem,
                        )
                        transcript_jobs.append(
                            {
                                "base_name": base_name,
                                "event_info": info,
                                "window": (window_start, window_end),
                            }
                        )
                        if idx == 0:
                            primary_event_info = info

                    if primary_event_info:
                        additional_count = len(matching_events) - 1
                        summary_suffix = f"and {additional_count} other(s)" if additional_count > 0 else ""
                        try:
                            rename_time = primary_event_info.get("start") or recording_start
                            working_path, _ = rename_audio_with_event(
                                working_path,
                                primary_event_info,
                                rename_time,
                                summary_suffix=summary_suffix,
                            )
                            if additional_count > 0:
                                print(
                                    f"Matched {len(matching_events)} calendar events "
                                    f"(primary: '{primary_event_info['summary']}'); renamed file to {working_path.name}"
                                )
                            else:
                                print(
                                    f"Matched calendar event '{primary_event_info['summary']}' "
                                    f"and renamed file to {working_path.name}"
                                )
                        except OSError as exc:
                            print(
                                f"Warning: failed to rename {working_path.name} ({exc}); continuing without rename.",
                                file=sys.stderr,
                            )
                elif fallback_event:
                    primary_event_info = attach_matching_calendars(
                        extract_event_metadata(fallback_event),
                        matching_calendars,
                    )
                    fallback_start = primary_event_info.get("start") or recording_start
                    fallback_end = primary_event_info.get("end") or fallback_start
                    window_start, window_end = compute_event_window(
                        recording_start,
                        recording_end,
                        fallback_start,
                        fallback_end,
                        duration_seconds,
                        EVENT_WINDOW_BUFFER_SECONDS,
                    )
                    base_name = build_event_base_name(
                        fallback_start,
                        primary_event_info.get("summary", "Untitled Event"),
                        source_stem,
                    )
                    transcript_jobs.append(
                        {
                            "base_name": base_name,
                            "event_info": primary_event_info,
                            "window": (window_start, window_end),
                        }
                    )
                    try:
                        rename_time = primary_event_info.get("start") or recording_start
                        working_path, _ = rename_audio_with_event(
                            working_path,
                            primary_event_info,
                            rename_time,
                        )
                        print(
                            f"Matched calendar event '{primary_event_info['summary']}' "
                            f"and renamed file to {working_path.name}"
                        )
                    except OSError as exc:
                        print(
                            f"Warning: failed to rename {working_path.name} ({exc}); continuing without rename.",
                            file=sys.stderr,
                        )
                else:
                    print("No calendar event found near the recording window; continuing without rename.")
            except Exception as exc:
                print(
                    f"Warning: calendar lookup failed ({exc}); continuing without calendar metadata.",
                    file=sys.stderr,
                )

    output_dir = ensure_output_dir(args.output_dir, working_path)

    if not transcript_jobs and primary_event_info:
        base_time = primary_event_info.get("start") or recording_start
        base_end = primary_event_info.get("end") or recording_end
        window_start, window_end = compute_event_window(
            recording_start,
            recording_end,
            base_time,
            base_end,
            duration_seconds,
            EVENT_WINDOW_BUFFER_SECONDS,
        )
        base_name = build_event_base_name(
            base_time,
            primary_event_info.get("summary", "Untitled Event"),
            source_stem,
        )
        transcript_jobs.append(
            {
                "base_name": base_name,
                "event_info": primary_event_info,
                "window": (window_start, window_end),
            }
        )

    if not transcript_jobs:
        transcript_jobs.append(
            {
                "base_name": working_path.stem,
                "event_info": None,
                "window": (0.0, duration_seconds),
            }
        )

    should_emit_docx = (not args.srt_output) or args.docx_output
    pending_artifacts: list[tuple[Path, TranscriptArtifact]] = []

    for job in transcript_jobs:
        base_name = job["base_name"]
        event_info = job.get("event_info")
        window_start, window_end = job.get("window") or (0.0, duration_seconds)
        selected_utterances = select_utterances_for_window(utterances, window_start, window_end)
        output_paths: dict[str, str] = {}

        if args.srt_output:
            speaker_names = {
                (utterance.get("speaker") or "").strip()
                for utterance in selected_utterances
                if (utterance.get("speaker") or "").strip()
            }
            suppress_speaker = len(speaker_names) <= 1
            srt_path = output_dir / f"{base_name} Transcript.srt"
            print(f"Writing SRT transcript to {srt_path}...", flush=True)
            write_srt(selected_utterances, srt_path, suppress_speaker=suppress_speaker)
            output_paths["srt"] = str(srt_path)

        if should_emit_docx:
            docx_path = output_dir / f"{base_name} Transcript.docx"
            print(f"Writing DOCX transcript to {docx_path}...", flush=True)
            write_docx(selected_utterances, docx_path, event_info=event_info, title=docx_title)
            output_paths["docx"] = str(docx_path)

        if args.text_output:
            text_path = output_dir / f"{base_name} Transcript.txt"
            print(f"Writing plain-text transcript to {text_path}...", flush=True)
            write_text(selected_utterances, text_path, event_info=event_info)
            output_paths["txt"] = str(text_path)

        artifact_path = output_dir / f"{base_name} Transcript.transcript.json"
        artifact = TranscriptArtifact(
            source_media_path=str(audio_path),
            working_media_path=str(working_path),
            backend=backend_name,
            duration_seconds=duration_seconds,
            recording_start=recording_start,
            recording_end=recording_end,
            transcript_window_start_seconds=window_start,
            transcript_window_end_seconds=window_end,
            utterance_count=len(selected_utterances),
            transcript_text="\n".join(format_utterance(utterance) for utterance in selected_utterances),
            utterances=selected_utterances,
            output_paths={**output_paths, "artifact": str(artifact_path)},
            event=event_info,
            transcript_title=docx_title,
        )
        pending_artifacts.append((artifact_path, artifact))

    embedded_media_path: Optional[Path] = None
    if args.embed_subtitles:
        speaker_names = {
            (utterance.get("speaker") or "").strip()
            for utterance in utterances
            if (utterance.get("speaker") or "").strip()
        }
        suppress_speaker = len(speaker_names) <= 1
        with tempfile.TemporaryDirectory() as tmp_dir:
            embed_srt_path = Path(tmp_dir) / "embedded.srt"
            write_srt(utterances, embed_srt_path, suppress_speaker=suppress_speaker)
            target_name = f"{working_path.stem} subtitled{working_path.suffix}"
            target_path = unique_path(working_path.with_name(target_name))
            try:
                print(f"Embedding subtitles into media file at {target_path}...", flush=True)
                embed_subtitles_with_ffmpeg(working_path, embed_srt_path, target_path)
                embedded_media_path = target_path
            except TranscriptionError as exc:
                print(f"Warning: failed to embed subtitles ({exc})", file=sys.stderr)

    for artifact_path, artifact in pending_artifacts:
        if embedded_media_path is not None:
            artifact.output_paths["embedded_media"] = str(embedded_media_path)
        print(f"Writing transcript artifact to {artifact_path}...", flush=True)
        artifact.write(artifact_path)
        if os.getenv("TRANSCRIPTS_STORE", "").strip().lower() in {"1", "true", "yes"}:
            try:
                ingest_artifact(
                    artifact_path,
                    root=Path(os.getenv("TRANSCRIPTS_STORE_DIR")).expanduser()
                    if os.getenv("TRANSCRIPTS_STORE_DIR")
                    else None,
                )
            except Exception as exc:
                print(f"Warning: failed to ingest transcript into ~/.transcripts ({exc})", file=sys.stderr)
        print(f"TRANSCRIPT_ARTIFACT_JSON={artifact_path}", flush=True)

    print("Completed successfully.")
    return True
