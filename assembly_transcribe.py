#!/usr/bin/env python3
"""
AssemblyAI transcription CLI for producing DOCX, TXT, or subtitle outputs with optional diarization.

Quick start:
1. python -m venv .venv
2. source .venv/bin/activate  (Windows: .venv\\Scripts\\activate)
3. pip install -r requirements.txt
4. export ASSEMBLYAI_API_KEY=your_key  (or copy api_keys.json.sample to api_keys.json)
5. python assembly_transcribe.py demo.wav --text-output

API keys are read in this order:
1. The `--api-key` argument.
2. The `ASSEMBLYAI_API_KEY` environment variable.
3. The `assemblyai_api_key` field inside a JSON file (default `api_keys.json`).

Language defaults to English (`en_us`) unless overridden via `assemblyai_language_code` in the same JSON file
(for example, `pt` for Portuguese).
"""
from __future__ import annotations

import argparse
import glob
import json
import os
import re
import sys
import subprocess
import time
import tempfile
import textwrap
import getpass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Generator, Iterable, Optional

import requests
from docx import Document
from docx.shared import Pt
from google.auth.exceptions import RefreshError
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build

DEFAULT_BASE_URL = "https://api.assemblyai.com"
DEFAULT_CHUNK_SIZE = 5 * 1024 * 1024  # 5 MB
DEFAULT_LANGUAGE_CODE = "en_us"
DEFAULT_OPENAI_MODEL = "gpt-4o-mini"
CALENDAR_SCOPES = ["https://www.googleapis.com/auth/calendar.readonly"]
SCRIPT_DIR = Path(__file__).resolve().parent
WILDCARD_PATTERN = re.compile(r"[*?\[\]]")
EVENT_WINDOW_BUFFER_SECONDS = 5 * 60
SENTENCE_BOUNDARY_RE = re.compile(r"(?<=[.!?])\s+")
MIN_SRT_CUE_DURATION = 0.5
CLI_DESCRIPTION = textwrap.dedent(
    """
    Upload one or more audio/video files to AssemblyAI, wait for completion, and export DOCX, TXT,
    or SRT outputs. Glob patterns are expanded by the script, and optional Google Calendar and
    ffmpeg integrations can enrich or embed the transcript.
    """
)
CLI_EPILOG = textwrap.dedent(
    """Examples:
      python assembly_transcribe.py meeting.m4a
      python assembly_transcribe.py meeting.m4a --text-output --output-dir transcripts
      python assembly_transcribe.py webinar.mp4 --srt-output
      python assembly_transcribe.py board_call.wav --use-calendar --embed-subtitles
    """
)


class AssemblyAIError(RuntimeError):
    """Raised when AssemblyAI returns an error payload."""

LANGUAGE_CODE_ALIASES = {
    "en": "en_us",
    "en_us": "en_us",
    "en-us": "en_us",
    "en_uk": "en_uk",
    "en-uk": "en_uk",
    "en_gb": "en_uk",
    "en-gb": "en_uk",
    "en_au": "en_au",
    "en-au": "en_au",
    "english": "en_us",
    "portuguese": "pt",
    "português": "pt",
    "portugues": "pt",
    "portugese": "pt",
    "pt": "pt",
    "pt_br": "pt",
    "pt-br": "pt",
    "pt_pt": "pt",
    "pt-pt": "pt",
}

def raise_for_status_with_details(response: requests.Response, *, context: str) -> None:
    try:
        response.raise_for_status()
    except requests.HTTPError as exc:
        detail: Optional[str] = None
        try:
            payload = response.json()
            if isinstance(payload, dict):
                detail = payload.get("error") or payload.get("message")
                if not detail:
                    detail = json.dumps(payload, ensure_ascii=False)
            else:
                detail = json.dumps(payload, ensure_ascii=False)
        except ValueError:
            text = (response.text or "").strip()
            if text:
                detail = text

        hint = ""
        if response.status_code == 400 and detail:
            lowered_detail = detail.lower()
            if not any(term in lowered_detail for term in ("balance", "top up", "quota", "billing", "payment")):
                hint = " (Hint: try `--no-speaker-labels` and/or a different `--model`.)"
        if detail:
            raise AssemblyAIError(f"{context} failed ({response.status_code}): {detail}{hint}") from exc
        raise AssemblyAIError(f"{context} failed ({response.status_code}).{hint}") from exc


def to_rfc3339(dt: datetime) -> str:
    return dt.astimezone(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def get_file_modified_time(path: Path) -> datetime:
    stats = path.stat()
    return datetime.fromtimestamp(stats.st_mtime, tz=timezone.utc)


def load_client_config(credentials_path: Path) -> dict[str, Any]:
    with credentials_path.open("r", encoding="utf-8") as handle:
        config = json.load(handle)
    if "installed" in config:
        return config
    if "web" in config:
        # Convert web client credentials into installed-style config accepted by InstalledAppFlow.
        web_config = config["web"]
        # Installed flow expects redirect URIs with localhost and urn scheme.
        installed_config = {
            "installed": {
                "client_id": web_config.get("client_id"),
                "project_id": web_config.get("project_id"),
                "auth_uri": web_config.get("auth_uri"),
                "token_uri": web_config.get("token_uri"),
                "auth_provider_x509_cert_url": web_config.get("auth_provider_x509_cert_url"),
                "client_secret": web_config.get("client_secret"),
                "redirect_uris": web_config.get(
                    "redirect_uris",
                    ["urn:ietf:wg:oauth:2.0:oob", "http://localhost"],
                ),
            }
        }
        return installed_config
    raise AssemblyAIError(
        f"Unsupported Google client secrets structure in {credentials_path}. Expected 'installed' or 'web'."
    )


def build_calendar_service(credentials_path: Path, token_path: Path, fallback_client_path: Optional[Path] = None):
    client_config: Optional[dict[str, Any]] = None
    chosen_config_path: Optional[Path] = None

    candidate_paths = [credentials_path]
    if fallback_client_path:
        candidate_paths.append(fallback_client_path)
    candidate_paths.extend(
        [
            credentials_path.with_name("client_secrets.json"),
            Path("client_secrets.json"),
            SCRIPT_DIR / "credentials.json",
            SCRIPT_DIR / "client_secrets.json",
        ]
    )

    unique_candidates: list[Path] = []
    seen_paths: set[Path] = set()
    for candidate in candidate_paths:
        if not candidate:
            continue
        resolved = candidate.expanduser().resolve()
        if resolved in seen_paths:
            continue
        seen_paths.add(resolved)
        unique_candidates.append(resolved)

    for candidate in unique_candidates:
        if not candidate.exists():
            continue
        try:
            client_config = load_client_config(candidate)
            chosen_config_path = candidate
            break
        except AssemblyAIError:
            continue

    if not client_config or not chosen_config_path:
        raise AssemblyAIError(
            "Google client secrets not found or invalid. Provide --calendar-credentials pointing to a "
            "client secret JSON downloaded from Google Cloud Console."
    )

    creds: Optional[Credentials] = None
    if token_path.exists():
        try:
            creds = Credentials.from_authorized_user_file(str(token_path), CALENDAR_SCOPES)
        except Exception:
            try:
                token_path.unlink(missing_ok=True)
            except OSError:
                pass
            creds = None

    credentials_updated = False
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
                credentials_updated = True
            except RefreshError:
                print("Google Calendar token expired or revoked; requesting new authorization.", file=sys.stderr)
                try:
                    token_path.unlink(missing_ok=True)
                except OSError:
                    pass
                creds = None
        if not creds or not creds.valid:
            flow = InstalledAppFlow.from_client_config(client_config, CALENDAR_SCOPES)
            creds = flow.run_local_server(port=0)
            credentials_updated = True
    if creds and credentials_updated:
        token_path.parent.mkdir(parents=True, exist_ok=True)
        token_path.write_text(creds.to_json(), encoding="utf-8")

    return build("calendar", "v3", credentials=creds, cache_discovery=False)


def parse_event_datetime(event_time: dict) -> Optional[datetime]:
    if not event_time:
        return None
    raw = event_time.get("dateTime") or event_time.get("date")
    if not raw:
        return None
    if len(raw) == 10:
        raw = f"{raw}T00:00:00+00:00"
    if raw.endswith("Z"):
        raw = raw[:-1] + "+00:00"
    if "T" in raw and "+" not in raw[10:] and "-" not in raw[10:]:
        raw = f"{raw}+00:00"
    return datetime.fromisoformat(raw)


def format_event_datetime(dt_value: Optional[datetime]) -> Optional[str]:
    if not dt_value:
        return None
    return dt_value.astimezone().strftime("%Y-%m-%d %H:%M %Z")


def extract_event_metadata(event: dict) -> dict[str, Any]:
    start_dt = parse_event_datetime(event.get("start", {}))
    end_dt = parse_event_datetime(event.get("end", {}))
    attendees = []
    for attendee in event.get("attendees", []):
        if attendee.get("responseStatus") == "declined":
            continue
        name = attendee.get("displayName")
        email = attendee.get("email")
        if name and email:
            attendees.append(f"{name} <{email}>")
        elif name:
            attendees.append(name)
        elif email:
            attendees.append(email)
    return {
        "summary": event.get("summary") or "Untitled Event",
        "start": start_dt,
        "end": end_dt,
        "location": event.get("location"),
        "attendees": attendees,
        "hangoutLink": event.get("hangoutLink"),
    }


def find_matching_events(
    service,
    calendar_id: str,
    recording_start: datetime,
    recording_end: datetime,
    window_hours: float,
) -> tuple[list[dict[str, Any]], Optional[dict]]:
    if recording_end < recording_start:
        recording_end = recording_start

    time_min = to_rfc3339(recording_start - timedelta(hours=window_hours))
    time_max = to_rfc3339(recording_end + timedelta(hours=window_hours))
    midpoint = recording_start + (recording_end - recording_start) / 2
    result = (
        service.events()
        .list(
            calendarId=calendar_id,
            timeMin=time_min,
            timeMax=time_max,
            singleEvents=True,
            maxResults=50,
            orderBy="startTime",
        )
        .execute()
    )
    events = result.get("items", [])
    if not events:
        return [], None

    best_event: Optional[dict] = None
    best_score: Optional[tuple] = None
    matching_events: list[dict[str, Any]] = []

    for event in events:
        event_start = parse_event_datetime(event.get("start", {}))
        event_end = parse_event_datetime(event.get("end", {}))
        if not event_start:
            continue

        event_range_end = event_end or event_start
        if event_range_end < event_start:
            event_range_end = event_start

        overlap_start = max(recording_start, event_start)
        overlap_end = min(recording_end, event_range_end)
        overlap_seconds = (overlap_end - overlap_start).total_seconds()
        if overlap_seconds < 0:
            overlap_seconds = 0.0

        event_duration_seconds = (event_range_end - event_start).total_seconds()
        if event_duration_seconds < 0:
            event_duration_seconds = 0.0
        coverage = overlap_seconds / event_duration_seconds if event_duration_seconds > 0 else 0.0

        if coverage >= 0.5:
            matching_events.append(
                {
                    "event": event,
                    "start": event_start,
                    "end": event_range_end,
                    "overlap_seconds": overlap_seconds,
                    "coverage": coverage,
                }
            )

        distances: list[float] = [abs((event_start - midpoint).total_seconds())]
        if event_end:
            distances.append(abs((event_end - midpoint).total_seconds()))
        distance = min(distances)

        if overlap_seconds > 0:
            score = (0, -overlap_seconds, distance)
        else:
            score = (1, distance)

        if best_score is None or score < best_score:
            best_score = score
            best_event = event

    matching_events.sort(key=lambda item: (item["start"], -item["overlap_seconds"], -item["coverage"]))
    return matching_events, best_event


def sanitize_filename_part(value: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]", " ", value)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def unique_path(base_path: Path) -> Path:
    counter = 1
    candidate = base_path
    while candidate.exists():
        candidate = base_path.with_name(f"{base_path.stem} ({counter}){base_path.suffix}")
        counter += 1
    return candidate


def build_event_base_name(recording_time: datetime, event_summary: str, source_stem: str) -> str:
    timestamp_str = recording_time.astimezone().strftime("%Y-%m-%d %H-%M")
    cleaned_summary = sanitize_filename_part(event_summary or "Untitled Event")
    cleaned_source = sanitize_filename_part(source_stem)
    parts = [timestamp_str, cleaned_summary, cleaned_source]
    return " ".join(part for part in parts if part)


def rename_audio_with_event(
    audio_path: Path,
    event_info: dict[str, Any],
    recording_time: datetime,
    *,
    summary_suffix: str = "",
) -> tuple[Path, str]:
    event_summary = event_info.get("summary", "Untitled Event")
    if summary_suffix:
        event_summary = f"{event_summary} {summary_suffix}"
    base_name = build_event_base_name(recording_time, event_summary, audio_path.stem)
    target_path = audio_path.with_name(f"{base_name}{audio_path.suffix}")
    target_path = unique_path(target_path)
    if target_path != audio_path:
        audio_path = audio_path.rename(target_path)
    return audio_path, base_name


def compute_event_window(
    recording_start: datetime,
    recording_end: datetime,
    event_start: Optional[datetime],
    event_end: Optional[datetime],
    duration_seconds: float,
    buffer_seconds: float,
) -> tuple[float, float]:
    event_start = event_start or event_end or recording_start
    event_end = event_end or event_start

    window_start_dt = max(recording_start, event_start - timedelta(seconds=buffer_seconds))
    window_end_dt = min(recording_end, event_end + timedelta(seconds=buffer_seconds))

    if window_end_dt < window_start_dt:
        window_end_dt = window_start_dt

    start_offset = (window_start_dt - recording_start).total_seconds()
    end_offset = (window_end_dt - recording_start).total_seconds()

    total_duration = max(duration_seconds, 0.0)
    start_offset = min(max(0.0, start_offset), total_duration)
    end_offset = min(max(0.0, end_offset), total_duration)
    if end_offset < start_offset:
        end_offset = start_offset

    return start_offset, end_offset


def select_utterances_for_window(
    utterances: list[dict],
    start_seconds: float,
    end_seconds: float,
) -> list[dict]:
    if end_seconds <= start_seconds:
        return [dict(utterance) for utterance in utterances]

    selected: list[dict] = []
    for utterance in utterances:
        utter_start = (utterance.get("start") or 0) / 1000
        utter_end = (utterance.get("end") or utter_start) / 1000
        if utter_end < start_seconds or utter_start > end_seconds:
            continue
        selected.append(dict(utterance))

    if not selected:
        return [dict(utterance) for utterance in utterances]

    return selected


def expand_audio_inputs(inputs: Iterable[str]) -> list[Path]:
    resolved: list[Path] = []
    seen: set[Path] = set()
    for raw in inputs:
        expanded = Path(raw).expanduser()
        matches: list[str] = []
        if WILDCARD_PATTERN.search(str(raw)):
            matches = glob.glob(str(expanded), recursive=True)
        else:
            if expanded.exists():
                matches = [str(expanded)]
            else:
                matches = glob.glob(str(expanded), recursive=True)

        if not matches:
            print(f"Warning: no files matched '{raw}'", file=sys.stderr)
            continue

        for match in matches:
            match_path = Path(match).expanduser().resolve()
            if not match_path.is_file():
                print(f"Skipping non-file match: {match_path}", file=sys.stderr)
                continue
            if match_path in seen:
                continue
            seen.add(match_path)
            resolved.append(match_path)

    return resolved

def parse_args(argv: Optional[Iterable[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=CLI_DESCRIPTION,
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=CLI_EPILOG,
    )
    parser.add_argument(
        "audio_inputs",
        nargs="+",
        help="Audio/video file(s) or glob patterns to transcribe (quote globs on Windows shells).",
    )

    api_key_group = parser.add_mutually_exclusive_group()
    api_key_group.add_argument(
        "--api-key",
        dest="api_key",
        help="AssemblyAI API key. Overrides environment variables and api_keys.json.",
    )
    api_key_group.add_argument(
        "--api-key-stdin",
        action="store_true",
        help="Read AssemblyAI API key from stdin (first line). Useful to avoid saving keys to disk.",
    )
    api_key_group.add_argument(
        "--api-key-prompt",
        action="store_true",
        help="Prompt for AssemblyAI API key interactively (input hidden).",
    )
    parser.add_argument(
        "--api-key-file",
        default="api_keys.json",
        help="Path to JSON file containing `assemblyai_api_key` (default: %(default)s). Copy api_keys.json.sample to create one.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        help="Directory for transcript outputs. Defaults to the audio file's directory.",
    )
    parser.add_argument(
        "--model",
        default="universal",
        help="AssemblyAI speech model (default: %(default)s).",
    )
    parser.add_argument(
        "--language",
        dest="language",
        help="Transcription language (overrides config). Examples: en_us (default), pt, en_uk, or auto for detection.",
    )
    parser.add_argument(
        "--poll-interval",
        type=float,
        default=3.0,
        help="Seconds between polling attempts (default: %(default)s).",
    )
    parser.add_argument(
        "--text-output",
        action="store_true",
        help="Also emit a plain-text transcript alongside the DOCX file.",
    )
    parser.add_argument(
        "--srt-output",
        action="store_true",
        help="Emit an SRT subtitle file instead of a DOCX transcript.",
    )
    parser.add_argument(
        "--docx-output",
        action="store_true",
        help="Also emit a DOCX transcript when --srt-output is set.",
    )
    parser.add_argument(
        "--embed-subtitles",
        action="store_true",
        help="Embed generated subtitles into the source media using ffmpeg.",
    )
    parser.add_argument(
        "--translate-to",
        dest="translate_to",
        help="Translate transcript and subtitles to a target language (e.g. en). Requires an OpenAI key.",
    )

    openai_key_group = parser.add_mutually_exclusive_group()
    openai_key_group.add_argument(
        "--openai-api-key",
        dest="openai_api_key",
        help="OpenAI API key (used for --translate-to). Overrides OPENAI_API_KEY and api_keys.json.",
    )
    openai_key_group.add_argument(
        "--openai-api-key-stdin",
        action="store_true",
        help="Read OpenAI API key from stdin (first line).",
    )
    openai_key_group.add_argument(
        "--openai-api-key-prompt",
        action="store_true",
        help="Prompt for OpenAI API key interactively (input hidden).",
    )
    parser.add_argument(
        "--openai-model",
        default=DEFAULT_OPENAI_MODEL,
        help="OpenAI model for translation (default: %(default)s).",
    )
    parser.add_argument(
        "--use-calendar",
        action="store_true",
        help="Match the audio file to a Google Calendar event and include metadata.",
    )
    parser.add_argument(
        "--calendar-id",
        default="primary",
        help="Google Calendar ID to query when --use-calendar is set (default: %(default)s).",
    )
    parser.add_argument(
        "--calendar-credentials",
        type=Path,
        default=SCRIPT_DIR / "credentials.json",
        help="Path to Google OAuth client credentials when using --use-calendar (default: %(default)s). "
        "If this file contains legacy tokens, the script will look for client secrets nearby.",
    )
    parser.add_argument(
        "--calendar-client-secrets",
        type=Path,
        default=SCRIPT_DIR / "client_secrets.json",
        help="Optional explicit path to Google client secrets (used when --calendar-credentials contains tokens).",
    )
    parser.add_argument(
        "--calendar-token",
        type=Path,
        default=SCRIPT_DIR / "token.json",
        help="Path to store Google OAuth tokens when using --use-calendar (default: %(default)s).",
    )
    parser.add_argument(
        "--calendar-window",
        type=float,
        default=24.0,
        help="Hours on either side of the file timestamp to search for matching events (default: %(default)s).",
    )
    speaker_group = parser.add_mutually_exclusive_group()
    speaker_group.add_argument(
        "--speaker-labels",
        dest="speaker_labels",
        action="store_true",
        help="Enable speaker diarization (default).",
    )
    speaker_group.add_argument(
        "--no-speaker-labels",
        dest="speaker_labels",
        action="store_false",
        help="Disable speaker diarization.",
    )
    parser.set_defaults(speaker_labels=True)
    return parser.parse_args(argv)


def resolve_config_candidates(configured_path: Path) -> list[Path]:
    if configured_path.is_absolute():
        return [configured_path]
    return [
        (Path.cwd() / configured_path).resolve(),
        (SCRIPT_DIR / configured_path.name).resolve(),
    ]


def load_json_config(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as fh:
        payload = json.load(fh)
    if not isinstance(payload, dict):
        raise AssemblyAIError(f"Invalid JSON structure in {path}; expected an object.")
    return payload


def normalize_language_code(raw_value: Optional[str]) -> tuple[Optional[str], bool]:
    if raw_value is None:
        return DEFAULT_LANGUAGE_CODE, False

    normalized = raw_value.strip()
    if not normalized:
        return DEFAULT_LANGUAGE_CODE, False

    lowered = normalized.lower()
    if lowered in {"auto", "detect", "language_detection", "language-detection"}:
        return None, True

    return LANGUAGE_CODE_ALIASES.get(lowered, normalized), False


def resolve_api_key(args: argparse.Namespace) -> str:
    if args.api_key:
        return args.api_key

    if getattr(args, "api_key_prompt", False):
        if not sys.stdin.isatty():
            raise AssemblyAIError("--api-key-prompt requires an interactive terminal (stdin is not a TTY).")
        entered = getpass.getpass("AssemblyAI API key: ").strip()
        if entered:
            return entered

    if getattr(args, "api_key_stdin", False):
        if sys.stdin.isatty():
            entered = sys.stdin.readline().strip()
        else:
            entered = sys.stdin.read().strip()
        if entered:
            return entered

    env_key = os.getenv("ASSEMBLYAI_API_KEY")
    if env_key:
        return env_key

    seen: set[Path] = set()
    for candidate in resolve_config_candidates(Path(args.api_key_file).expanduser()):
        if candidate in seen:
            continue
        seen.add(candidate)
        if not candidate.exists():
            continue
        try:
            payload = load_json_config(candidate)
        except json.JSONDecodeError as exc:
            raise AssemblyAIError(f"Invalid JSON in {candidate}: {exc}") from exc

        for key_field in ("assemblyai_api_key", "assembly_ai_api_key"):
            api_key = payload.get(key_field)
            if api_key:
                return api_key

    raise AssemblyAIError(
        "AssemblyAI API key not found. Provide --api-key, set ASSEMBLYAI_API_KEY, "
        f"or store it in {args.api_key_file} (or alongside the script) under 'assemblyai_api_key'."
    )


def resolve_language_settings(args: argparse.Namespace) -> tuple[Optional[str], bool]:
    if getattr(args, "language", None):
        return normalize_language_code(args.language)

    for candidate in resolve_config_candidates(Path(args.api_key_file).expanduser()):
        if not candidate.exists():
            continue
        try:
            payload = load_json_config(candidate)
        except json.JSONDecodeError as exc:
            print(f"Warning: skipping invalid JSON in {candidate} ({exc}); using English defaults.", file=sys.stderr)
            return DEFAULT_LANGUAGE_CODE, False
        except AssemblyAIError as exc:
            print(f"Warning: skipping {candidate} ({exc}); using English defaults.", file=sys.stderr)
            return DEFAULT_LANGUAGE_CODE, False

        for field in ("assemblyai_language_code", "assembly_ai_language_code"):
            if field in payload:
                return normalize_language_code(payload.get(field))
        break

    return DEFAULT_LANGUAGE_CODE, False


def resolve_openai_key(args: argparse.Namespace) -> Optional[str]:
    direct_key = getattr(args, "openai_api_key", None)
    if direct_key:
        return direct_key

    if getattr(args, "openai_api_key_prompt", False):
        if not sys.stdin.isatty():
            raise AssemblyAIError("--openai-api-key-prompt requires an interactive terminal (stdin is not a TTY).")
        entered = getpass.getpass("OpenAI API key: ").strip()
        if entered:
            return entered

    if getattr(args, "openai_api_key_stdin", False):
        if sys.stdin.isatty():
            entered = sys.stdin.readline().strip()
        else:
            entered = sys.stdin.read().strip()
        if entered:
            return entered

    env_key = os.getenv("OPENAI_API_KEY")
    if env_key:
        return env_key

    if not getattr(args, "api_key_file", None):
        return None

    for candidate in resolve_config_candidates(Path(args.api_key_file).expanduser()):
        if not candidate.exists():
            continue
        try:
            payload = load_json_config(candidate)
        except json.JSONDecodeError:
            return None
        except AssemblyAIError:
            return None

        for key_field in ("openai_api_key", "open_ai_api_key"):
            openai_key = payload.get(key_field)
            if openai_key:
                return openai_key
        break

    return None


def stream_file(path: Path, chunk_size: int = DEFAULT_CHUNK_SIZE) -> Generator[bytes, None, None]:
    with path.open("rb") as handle:
        while True:
            chunk = handle.read(chunk_size)
            if not chunk:
                break
            yield chunk


def upload_audio(session: requests.Session, audio_path: Path) -> str:
    upload_endpoint = f"{DEFAULT_BASE_URL}/v2/upload"
    response = session.post(upload_endpoint, data=stream_file(audio_path))
    raise_for_status_with_details(response, context="Upload")
    data = response.json()
    if "upload_url" not in data:
        raise AssemblyAIError(f"Upload response missing 'upload_url': {data}")
    return data["upload_url"]


def request_transcription(
    session: requests.Session,
    audio_url: str,
    *,
    model: str,
    speaker_labels: bool,
    language_code: Optional[str],
    language_detection: bool,
) -> str:
    transcript_endpoint = f"{DEFAULT_BASE_URL}/v2/transcript"
    payload = {
        "audio_url": audio_url,
        "speech_model": model,
        "speaker_labels": speaker_labels,
    }
    if language_detection:
        payload["language_detection"] = True
    else:
        payload["language_code"] = language_code or DEFAULT_LANGUAGE_CODE
    response = session.post(transcript_endpoint, json=payload)
    raise_for_status_with_details(response, context="Transcription request")
    data = response.json()
    if "id" not in data:
        raise AssemblyAIError(f"Transcript response missing 'id': {data}")
    return data["id"]


def poll_transcript(
    session: requests.Session,
    transcript_id: str,
    *,
    poll_interval: float,
) -> dict:
    status_endpoint = f"{DEFAULT_BASE_URL}/v2/transcript/{transcript_id}"
    while True:
        response = session.get(status_endpoint)
        raise_for_status_with_details(response, context="Polling transcript")
        payload = response.json()

        status = payload.get("status")
        if status == "completed":
            return payload
        if status == "error":
            raise AssemblyAIError(payload.get("error", "Unknown AssemblyAI error"))

        ETA = payload.get("eta")
        if ETA is not None:
            print(f"Waiting... status={status}, eta={ETA}s", flush=True)
        else:
            print(f"Waiting... status={status}", flush=True)
        time.sleep(poll_interval)


def format_utterance(utterance: dict) -> str:
    speaker = utterance.get("speaker", "Speaker")
    start = utterance.get("start", 0) / 1000
    end = utterance.get("end", 0) / 1000
    text = (utterance.get("text") or "").strip()
    return f"{speaker} [{start:0.2f}s - {end:0.2f}s]: {text}"


def add_event_metadata_to_doc(document: Document, event_info: dict[str, Any]) -> None:
    document.add_heading("Event Details", level=2)
    document.add_paragraph(f"Event: {event_info['summary']}")

    start_str = format_event_datetime(event_info.get("start"))
    end_str = format_event_datetime(event_info.get("end"))
    if start_str:
        document.add_paragraph(f"Start: {start_str}")
    if end_str:
        document.add_paragraph(f"End: {end_str}")
    if event_info.get("location"):
        document.add_paragraph(f"Location: {event_info['location']}")
    if event_info.get("hangoutLink"):
        document.add_paragraph(f"Meeting Link: {event_info['hangoutLink']}")

    attendees = event_info.get("attendees") or []
    if attendees:
        document.add_paragraph("Attendees:")
        for attendee in attendees:
            document.add_paragraph(attendee, style="List Bullet")
    document.add_paragraph("")


def write_docx(
    utterances: list[dict],
    output_path: Path,
    *,
    event_info: Optional[dict[str, Any]] = None,
) -> None:
    document = Document()
    document.add_heading("AssemblyAI Transcript", level=1)
    if event_info:
        add_event_metadata_to_doc(document, event_info)
    for utterance in utterances:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(format_utterance(utterance))
        run.font.size = Pt(11)
    document.save(output_path)


def split_into_sentences(text: str) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []
    sentences = SENTENCE_BOUNDARY_RE.split(text)
    return [sentence.strip() for sentence in sentences if sentence.strip()]


def openai_translate_lines(
    *,
    api_key: str,
    model: str,
    target_language: str,
    lines: list[str],
    timeout: float = 90.0,
) -> list[str]:
    endpoint = "https://api.openai.com/v1/chat/completions"
    session = requests.Session()
    session.headers.update({"authorization": f"Bearer {api_key}"})

    system_prompt = (
        "You are a translation engine. Translate the provided JSON array of strings into "
        f"{target_language}. Output ONLY a JSON array of strings of the same length, in the same order. "
        "Do not add commentary, numbering, or extra keys."
    )
    payload = {
        "model": model,
        "temperature": 0,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps(lines, ensure_ascii=False)},
        ],
    }
    response = session.post(endpoint, json=payload, timeout=timeout)
    raise_for_status_with_details(response, context="OpenAI translation")
    data = response.json()
    content = (((data.get("choices") or [{}])[0]).get("message") or {}).get("content") or ""
    content = content.strip()
    if not content:
        raise AssemblyAIError("OpenAI translation returned an empty response.")

    try:
        translated = json.loads(content)
    except json.JSONDecodeError as exc:
        raise AssemblyAIError("OpenAI translation did not return valid JSON.") from exc

    if not isinstance(translated, list) or not all(isinstance(item, str) for item in translated):
        raise AssemblyAIError("OpenAI translation returned unexpected JSON (expected a string array).")
    if len(translated) != len(lines):
        raise AssemblyAIError(f"OpenAI translation returned {len(translated)} lines but expected {len(lines)}.")
    return [item.strip() for item in translated]


def translate_utterances_openai(
    utterances: list[dict],
    *,
    api_key: str,
    model: str,
    target_language: str,
) -> list[dict]:
    translated_utterances = [dict(utterance) for utterance in utterances]
    sentence_map: list[list[str]] = []
    all_sentences: list[str] = []

    for utterance in translated_utterances:
        sentences = split_into_sentences(utterance.get("text") or "")
        if not sentences:
            sentence_map.append([])
            continue
        sentence_map.append(sentences)
        all_sentences.extend(sentences)

    if not all_sentences:
        return translated_utterances

    translated_sentences: list[str] = []
    batch_size = 25
    for start in range(0, len(all_sentences), batch_size):
        batch = all_sentences[start : start + batch_size]
        translated_sentences.extend(
            openai_translate_lines(
                api_key=api_key,
                model=model,
                target_language=target_language,
                lines=batch,
            )
        )

    idx = 0
    for utterance, original_sentences in zip(translated_utterances, sentence_map):
        if not original_sentences:
            utterance.pop("sentences", None)
            continue
        translated_for_utterance = translated_sentences[idx : idx + len(original_sentences)]
        idx += len(original_sentences)
        utterance["sentences"] = translated_for_utterance
        utterance["text"] = " ".join(translated_for_utterance).strip()

    return translated_utterances


def format_srt_timestamp(seconds: float) -> str:
    seconds = max(0.0, seconds)
    total_ms = int(round(seconds * 1000))
    hours, remainder = divmod(total_ms, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, millis = divmod(remainder, 1_000)
    return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"


def allocate_sentence_timings(
    start: float,
    end: float,
    sentences: list[str],
) -> list[tuple[float, float]]:
    if not sentences:
        return []
    # Ensure timing window is non-negative.
    start = max(0.0, start)
    duration = max(0.0, end - start)
    if duration <= 0.0:
        duration = MIN_SRT_CUE_DURATION * len(sentences)
        end = start + duration
    else:
        end = start + duration

    weights = [max(len(sentence.split()), 1) for sentence in sentences]
    total_weight = sum(weights) or len(sentences)

    timings: list[tuple[float, float]] = []
    current_start = start
    for idx, weight in enumerate(weights):
        if idx == len(weights) - 1:
            cue_end = max(current_start + MIN_SRT_CUE_DURATION, end)
        else:
            share = weight / total_weight
            cue_duration = duration * share
            cue_end = current_start + cue_duration
            if cue_end - current_start < MIN_SRT_CUE_DURATION:
                cue_end = current_start + MIN_SRT_CUE_DURATION
        timings.append((current_start, cue_end))
        current_start = cue_end
    return timings


def write_srt(
    utterances: list[dict],
    output_path: Path,
    *,
    suppress_speaker: bool,
) -> None:
    cues: list[tuple[int, float, float, str]] = []
    cue_index = 1
    previous_end = 0.0
    for utterance in utterances:
        start = (utterance.get("start") or 0) / 1000.0
        end = (utterance.get("end") or 0) / 1000.0
        if start < previous_end:
            start = previous_end
        if end < start:
            end = start
        sentences = utterance.get("sentences")
        if not isinstance(sentences, list) or not all(isinstance(item, str) for item in sentences):
            sentences = split_into_sentences(utterance.get("text") or "")
        if not sentences:
            continue
        timings = allocate_sentence_timings(start, end, sentences)
        speaker = (utterance.get("speaker") or "").strip() or "Speaker"
        for sentence, (sentence_start, sentence_end) in zip(sentences, timings):
            if sentence_end <= sentence_start:
                sentence_end = sentence_start + MIN_SRT_CUE_DURATION
            text = sentence
            if not suppress_speaker:
                text = f"{speaker}: {sentence}"
            cues.append((cue_index, sentence_start, sentence_end, text))
            cue_index += 1
            previous_end = max(previous_end, sentence_end)

    with output_path.open("w", encoding="utf-8") as handle:
        for idx, start_ts, end_ts, text in cues:
            handle.write(f"{idx}\n")
            handle.write(f"{format_srt_timestamp(start_ts)} --> {format_srt_timestamp(end_ts)}\n")
            handle.write(f"{text}\n\n")


def determine_subtitle_codec(media_suffix: str) -> str:
    suffix = media_suffix.lower()
    if suffix in {".mp4", ".m4v", ".mov"}:
        return "mov_text"
    if suffix in {".mkv"}:
        return "srt"
    raise AssemblyAIError(
        f"Embedding subtitles is only supported for MP4/M4V/MOV/MKV files; unsupported suffix '{media_suffix}'."
    )


def embed_subtitles_with_ffmpeg(
    source_media: Path,
    subtitle_file: Path,
    output_media: Path,
) -> None:
    if not source_media.exists():
        raise AssemblyAIError(f"Source media file not found: {source_media}")
    if not subtitle_file.exists():
        raise AssemblyAIError(f"Subtitle file not found: {subtitle_file}")

    subtitle_codec = determine_subtitle_codec(source_media.suffix)
    command = [
        "ffmpeg",
        "-y",
        "-i",
        str(source_media),
        "-i",
        str(subtitle_file),
        "-c",
        "copy",
        "-c:s",
        subtitle_codec,
        "-map",
        "0",
        "-map",
        "1",
        str(output_media),
    ]
    try:
        completed = subprocess.run(command, check=False, capture_output=True, text=True)
    except FileNotFoundError as exc:
        raise AssemblyAIError("ffmpeg executable not found on PATH; install ffmpeg to embed subtitles.") from exc

    if completed.returncode != 0:
        stderr = (completed.stderr or "").strip()
        raise AssemblyAIError(f"ffmpeg failed to embed subtitles ({stderr or 'no error output provided'})")


def write_text(
    utterances: list[dict],
    output_path: Path,
    *,
    event_info: Optional[dict[str, Any]] = None,
) -> None:
    with output_path.open("w", encoding="utf-8") as handle:
        if event_info:
            handle.write(f"Event: {event_info['summary']}\n")
            start_str = format_event_datetime(event_info.get("start"))
            end_str = format_event_datetime(event_info.get("end"))
            if start_str:
                handle.write(f"Start: {start_str}\n")
            if end_str:
                handle.write(f"End: {end_str}\n")
            if event_info.get("location"):
                handle.write(f"Location: {event_info['location']}\n")
            if event_info.get("hangoutLink"):
                handle.write(f"Meeting Link: {event_info['hangoutLink']}\n")
            attendees = event_info.get("attendees") or []
            if attendees:
                handle.write("Attendees:\n")
                for attendee in attendees:
                    handle.write(f" - {attendee}\n")
            handle.write("\n")
        for utterance in utterances:
            handle.write(format_utterance(utterance))
            handle.write("\n")


def ensure_output_dir(path: Optional[Path], audio_path: Path) -> Path:
    if path is None:
        return audio_path.parent
    path.mkdir(parents=True, exist_ok=True)
    return path


def process_audio_file(
    audio_path: Path,
    args: argparse.Namespace,
    api_key: str,
    language_settings: tuple[Optional[str], bool],
    calendar_service,
) -> bool:
    working_path = audio_path
    transcript_jobs: list[dict[str, Any]] = []

    session = requests.Session()
    session.headers.update({"authorization": api_key, "user-agent": "assembly-transcribe-cli"})

    print(f"Uploading {working_path} to AssemblyAI...", flush=True)
    try:
        audio_url = upload_audio(session, working_path)
    except (requests.RequestException, AssemblyAIError) as exc:
        print(f"Upload failed: {exc}", file=sys.stderr)
        return False

    print("Starting transcription job...", flush=True)
    try:
        language_code, language_detection = language_settings
        transcript_id = request_transcription(
            session,
            audio_url,
            model=args.model,
            speaker_labels=args.speaker_labels,
            language_code=language_code,
            language_detection=language_detection,
        )
    except (requests.RequestException, AssemblyAIError) as exc:
        print(f"Transcription request failed: {exc}", file=sys.stderr)
        return False

    print(f"Polling transcript {transcript_id}...", flush=True)
    try:
        transcript_payload = poll_transcript(
            session,
            transcript_id,
            poll_interval=args.poll_interval,
        )
    except (requests.RequestException, AssemblyAIError) as exc:
        print(f"Polling failed: {exc}", file=sys.stderr)
        return False

    utterances = transcript_payload.get("utterances") or []
    if not utterances:
        print("AssemblyAI response did not contain diarized utterances; falling back to plain text.")
        text = transcript_payload.get("text") or ""
        utterances = [{"speaker": "Speaker", "start": 0, "end": 0, "text": text}]

    if args.translate_to:
        openai_key = resolve_openai_key(args)
        if not openai_key:
            print(
                "Error: --translate-to requires an OpenAI key. Provide --openai-api-key, set OPENAI_API_KEY, "
                "or add openai_api_key to api_keys.json.",
                file=sys.stderr,
            )
            return False
        print(f"Translating transcript to {args.translate_to}...", flush=True)
        try:
            utterances = translate_utterances_openai(
                utterances,
                api_key=openai_key,
                model=args.openai_model,
                target_language=args.translate_to,
            )
        except (requests.RequestException, AssemblyAIError) as exc:
            print(f"Translation failed: {exc}", file=sys.stderr)
            return False

    audio_duration = transcript_payload.get("audio_duration")
    if not audio_duration:
        max_end_ms = max((utterance.get("end") or 0) for utterance in utterances) if utterances else 0
        audio_duration = max_end_ms / 1000 if max_end_ms else 0

    try:
        duration_seconds = float(audio_duration)
    except (TypeError, ValueError):
        duration_seconds = 0.0
    if duration_seconds < 0:
        duration_seconds = 0.0

    recording_end = get_file_modified_time(working_path)
    recording_start = recording_end - timedelta(seconds=duration_seconds)
    if recording_start > recording_end:
        recording_start = recording_end

    source_stem = audio_path.stem
    primary_event_info: Optional[dict[str, Any]] = None

    if args.use_calendar:
        if calendar_service is None:
            print("Warning: calendar service unavailable; skipping event lookup.", file=sys.stderr)
        else:
            try:
                matching_events, fallback_event = find_matching_events(
                    calendar_service,
                    args.calendar_id,
                    recording_start,
                    recording_end,
                    args.calendar_window,
                )

                if matching_events:
                    for idx, match in enumerate(matching_events):
                        info = extract_event_metadata(match["event"])
                        event_start = match.get("start") or recording_start
                        event_end = match.get("end") or event_start
                        window_start, window_end = compute_event_window(
                            recording_start,
                            recording_end,
                            event_start,
                            event_end,
                            duration_seconds,
                            EVENT_WINDOW_BUFFER_SECONDS,
                        )
                        base_name = build_event_base_name(
                            event_start,
                            info.get("summary", "Untitled Event"),
                            source_stem,
                        )
                        transcript_jobs.append(
                            {
                                "base_name": base_name,
                                "event_info": info,
                                "window": (window_start, window_end),
                            }
                        )
                        if idx == 0:
                            primary_event_info = info

                    if primary_event_info:
                        additional_count = len(matching_events) - 1
                        summary_suffix = ""
                        if additional_count > 0:
                            summary_suffix = f"and {additional_count} other(s)"
                        try:
                            rename_time = primary_event_info.get("start") or recording_start
                            working_path, _ = rename_audio_with_event(
                                working_path,
                                primary_event_info,
                                rename_time,
                                summary_suffix=summary_suffix,
                            )
                            if additional_count > 0:
                                print(
                                    f"Matched {len(matching_events)} calendar events "
                                    f"(primary: '{primary_event_info['summary']}'); renamed file to {working_path.name}"
                                )
                            else:
                                print(
                                    f"Matched calendar event '{primary_event_info['summary']}' "
                                    f"and renamed file to {working_path.name}"
                                )
                        except OSError as exc:
                            print(
                                f"Warning: failed to rename {working_path.name} ({exc}); continuing without rename.",
                                file=sys.stderr,
                            )
                else:
                    if fallback_event:
                        primary_event_info = extract_event_metadata(fallback_event)
                        fallback_start = primary_event_info.get("start") or recording_start
                        fallback_end = primary_event_info.get("end") or fallback_start
                        window_start, window_end = compute_event_window(
                            recording_start,
                            recording_end,
                            fallback_start,
                            fallback_end,
                            duration_seconds,
                            EVENT_WINDOW_BUFFER_SECONDS,
                        )
                        base_name = build_event_base_name(
                            fallback_start,
                            primary_event_info.get("summary", "Untitled Event"),
                            source_stem,
                        )
                        transcript_jobs.append(
                            {
                                "base_name": base_name,
                                "event_info": primary_event_info,
                                "window": (window_start, window_end),
                            }
                        )
                        try:
                            rename_time = primary_event_info.get("start") or recording_start
                            working_path, _ = rename_audio_with_event(
                                working_path,
                                primary_event_info,
                                rename_time,
                            )
                            print(
                                f"Matched calendar event '{primary_event_info['summary']}' "
                                f"and renamed file to {working_path.name}"
                            )
                        except OSError as exc:
                            print(
                                f"Warning: failed to rename {working_path.name} ({exc}); continuing without rename.",
                                file=sys.stderr,
                            )
                    else:
                        print("No calendar event found near the recording window; continuing without rename.")
            except Exception as exc:
                print(
                    f"Warning: calendar lookup failed ({exc}); continuing without calendar metadata.",
                    file=sys.stderr,
                )

    output_dir = ensure_output_dir(args.output_dir, working_path)

    if not transcript_jobs and primary_event_info:
        base_time = primary_event_info.get("start") or recording_start
        base_end = primary_event_info.get("end") or recording_end
        window_start, window_end = compute_event_window(
            recording_start,
            recording_end,
            base_time,
            base_end,
            duration_seconds,
            EVENT_WINDOW_BUFFER_SECONDS,
        )
        base_name = build_event_base_name(
            base_time,
            primary_event_info.get("summary", "Untitled Event"),
            source_stem,
        )
        transcript_jobs.append(
            {
                "base_name": base_name,
                "event_info": primary_event_info,
                "window": (window_start, window_end),
            }
        )

    if not transcript_jobs:
        transcript_jobs.append(
            {
                "base_name": working_path.stem,
                "event_info": None,
                "window": (0.0, duration_seconds),
            }
        )

    should_emit_docx = (not args.srt_output) or args.docx_output
    should_emit_srt_files = args.srt_output

    for job in transcript_jobs:
        base_name = job["base_name"]
        event_info = job.get("event_info")
        window = job.get("window") or (0.0, duration_seconds)
        window_start, window_end = window
        selected_utterances = select_utterances_for_window(utterances, window_start, window_end)

        if should_emit_srt_files:
            speaker_names = {
                (utterance.get("speaker") or "").strip()
                for utterance in selected_utterances
                if (utterance.get("speaker") or "").strip()
            }
            suppress_speaker = len(speaker_names) <= 1
            srt_path = output_dir / f"{base_name} Transcript.srt"
            print(f"Writing SRT transcript to {srt_path}...", flush=True)
            write_srt(selected_utterances, srt_path, suppress_speaker=suppress_speaker)
        if should_emit_docx:
            docx_path = output_dir / f"{base_name} Transcript.docx"
            print(f"Writing DOCX transcript to {docx_path}...", flush=True)
            write_docx(selected_utterances, docx_path, event_info=event_info)

        if args.text_output:
            text_path = output_dir / f"{base_name} Transcript.txt"
            print(f"Writing plain-text transcript to {text_path}...", flush=True)
            write_text(selected_utterances, text_path, event_info=event_info)

    if args.embed_subtitles:
        speaker_names = {
            (utterance.get("speaker") or "").strip()
            for utterance in utterances
            if (utterance.get("speaker") or "").strip()
        }
        suppress_speaker = len(speaker_names) <= 1
        with tempfile.TemporaryDirectory() as tmp_dir:
            embed_srt_path = Path(tmp_dir) / "embedded.srt"
            write_srt(utterances, embed_srt_path, suppress_speaker=suppress_speaker)
            target_name = f"{working_path.stem} subtitled{working_path.suffix}"
            target_path = unique_path(working_path.with_name(target_name))
            try:
                print(f"Embedding subtitles into media file at {target_path}...", flush=True)
                embed_subtitles_with_ffmpeg(working_path, embed_srt_path, target_path)
            except AssemblyAIError as exc:
                print(f"Warning: failed to embed subtitles ({exc})", file=sys.stderr)

    print("Completed successfully.")
    return True


def main(argv: Optional[Iterable[str]] = None) -> int:
    args = parse_args(argv)
    audio_paths = expand_audio_inputs(args.audio_inputs)

    if not audio_paths:
        print("Error: no audio files matched the provided inputs.", file=sys.stderr)
        return 1

    try:
        api_key = resolve_api_key(args)
    except AssemblyAIError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    language_settings = resolve_language_settings(args)

    calendar_service = None
    if args.use_calendar:
        try:
            calendar_service = build_calendar_service(
                args.calendar_credentials,
                args.calendar_token,
                fallback_client_path=args.calendar_client_secrets,
            )
        except Exception as exc:
            print(
                f"Warning: calendar service initialization failed ({exc}); continuing without calendar metadata.",
                file=sys.stderr,
            )
            calendar_service = None

    any_failures = False
    for path in audio_paths:
        success = process_audio_file(path, args, api_key, language_settings, calendar_service)
        if not success:
            any_failures = True

    return 0 if not any_failures else 1


if __name__ == "__main__":
    sys.exit(main())
